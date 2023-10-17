#-*- conding:utf-8 -*-
import json
import requests
import pandas as pd
import time
import numpy as np
import os
import re
import datetime
from requests import Request,Session
from matplotlib import font_manager as fm, rcParams
from requests.exceptions import ConnectionError,Timeout,TooManyRedirects
import matplotlib.pyplot as plt
import seaborn as sns
import hashlib
import random
from PIL import Image, ImageDraw, ImageFont
import cv2
from watermarker.marker import add_mark
import telegram
def jigsaw(imgs, direction, gap=0):
    imgs = [Image.fromarray(img) for img in imgs]
    w, h = imgs[0].size
    if direction == "horizontal":
        result = Image.new(imgs[0].mode, ((w+gap)*len(imgs)-gap, h))
        for i, img in enumerate(imgs):
            result.paste(img, box=((w+gap)*i, 0))
    elif direction == "vertical":
        result = Image.new(imgs[0].mode, (w, (h+gap)*len(imgs)-gap))
        for i, img in enumerate(imgs):
            result.paste(img, box=(0, (h+gap)*i))
    else:
        raise ValueError("The direction parameter has only two options: horizontal and vertical")
    return np.array(result)
# ======= 正式开始执行
prop = fm.FontProperties(fname='/root/xianhuo/SimHei.ttf')
#pip install jojo-office
#import office
#office.open_file("output.docx").save("output.pdf")
date_now = datetime.datetime.utcnow()
# ----- 基础btc数据
url_address = [ 'https://api.glassnode.com/v1/metrics/market/price_usd_ohlc']
url_name = ['k_fold']
# insert your API key here
API_KEY = '26BLocpWTcSU7sgqDdKzMHMpJDm'
data_list = []
for num in range(len(url_name)):
    print(num)
    addr = url_address[num]
    name = url_name[num]
    # make API request
    res_addr = requests.get(addr,params={'a': 'BTC', 'api_key': API_KEY})
    # convert to pandas dataframe
    ins = pd.read_json(res_addr.text, convert_dates=['t'])
    #ins.to_csv('test.csv')
    #print(ins['o'])
    ins['date'] =  ins['t']
    ins['value'] =  ins['o']
    ins = ins[['date','value']]
    data_list.append(ins)
result_data = data_list[0][['date']]
for i in range(len(data_list)):
    df = data_list[i]
    result_data = result_data.merge(df,how='left',on='date')
#last_data = result_data[(result_data.date>='2016-01-01') & (result_data.date<='2020-01-01')]
last_data = result_data[(result_data.date>='2013-01-01')]
last_data = last_data.sort_values(by=['date'])
last_data = last_data.reset_index(drop=True)
print(type(last_data))
date = []
open_p = []
close_p = []
high_p = []
low_p = []
for i in range(len(last_data)):
    date.append(last_data['date'][i])
    open_p.append(last_data['value'][i]['o'])
    close_p.append(last_data['value'][i]['c'])
    high_p.append(last_data['value'][i]['h'])
    low_p.append(last_data['value'][i]['l'])
#res_data = pd.DataFrame({'date':date,'open':open_p,'close':close_p,'high':high_p,'low':low_p})
price_data = pd.DataFrame({'date':date,'close':close_p})
price_data = price_data.sort_values(by=['date'])
price_data = price_data.reset_index(drop=True)

url_address = ['https://api.glassnode.com/v1/metrics/market/price_usd_close']
url_name = ['Price']
# insert your API key here
API_KEY = '26BLocpWTcSU7sgqDdKzMHMpJDm'
data_list = []
for num in range(len(url_name)):
    print(num)
    addr = url_address[num]
    name = url_name[num]
    # make API request
    res_addr = requests.get(addr,params={'a': 'BTC', 'api_key': API_KEY})
    # convert to pandas dataframe
    ins = pd.read_json(res_addr.text, convert_dates=['t'])
    ins['date'] =  ins['t']
    ins[name] =  ins['v']
    ins = ins[['date',name]]
    data_list.append(ins)
result_data = data_list[0][['date']]
for i in range(len(data_list)):
    df = data_list[i]
    result_data = result_data.merge(df,how='left',on='date')
#last_data = result_data[(result_data.date>='2016-01-01') & (result_data.date<='2020-01-01')]
last_data = result_data[(result_data.date>='2010-01-01')]
from dateutil.relativedelta import relativedelta 
#last_data['new_date'] = last_data['date'].apply(lambda x:x + relativedelta(years=1))
last_data = last_data.sort_values(by=['date'])
last_data = last_data.reset_index(drop=True)
date = []
price_raw = []
price_ma120 = []
price_ma200 = []
price_ma4y = []
for j in range(len(last_data)-1824):
    ins = last_data[j:j+1825]
    ins = ins.sort_values(by='date')
    ins = ins.reset_index(drop=True)
    date.append(ins['date'][1824])
    price_raw.append(ins['Price'][1824])
    price_ma4y.append(np.mean(ins['Price'][-1459:]))
    price_ma200.append(np.mean(ins['Price'][-199:]))
    price_ma120.append(np.mean(ins['Price'][-119:]))
jun_df = pd.DataFrame({'date':date,'price_raw':price_raw,'price_ma120':price_ma120,'price_ma200':price_ma200,'price_ma4y':price_ma4y})
jun_df = jun_df[(jun_df.date>='2018-12-01')]
f, axes = plt.subplots(figsize=(20, 10))
axes_fu = axes.twinx()

sns.lineplot(x="date", y="price_raw", data=jun_df, color = 'black', linewidth=0.5,ci=95, ax=axes)
sns.lineplot(x="date", y="price_ma120",color='red', linewidth=0.5,data=jun_df, ci=95, ax=axes)
sns.lineplot(x="date", y="price_ma200",color='green', linewidth=0.5,data=jun_df, ci=95, ax=axes)
sns.lineplot(x="date", y="price_ma4y",color='blue', linewidth=0.5,data=jun_df, ci=95, ax=axes)
plt.title('BTC价格-120D/200D/4Y线', fontsize=20,fontproperties=prop) 
axes.set_xlabel('时间',fontsize=14,fontproperties=prop)
axes.set_ylabel("BTC价格",fontsize=14,fontproperties=prop)
plt.savefig('BTC价格120D.png',  bbox_inches='tight')
plt.close()

#add_mark(file = "BTC价格120D.png", out = "out",mark = "MMC研究猿卡森出品", opacity=0.2, angle=30, space=30)

# =======================================================================btc链上数据===================================================================

# ----- btc市值占比
url_address = [ 'https://api.glassnode.com/v1/metrics/market/btc_dominance']
url_name = ['k_fold']
# insert your API key here
API_KEY = '26BLocpWTcSU7sgqDdKzMHMpJDm'
data_list = []
for num in range(len(url_name)):
    print(num)
    addr = url_address[num]
    name = url_name[num]
    # make API request
    res_addr = requests.get(addr,params={'a': 'BTC', 'api_key': API_KEY})
    # convert to pandas dataframe
    ins = pd.read_json(res_addr.text, convert_dates=['t'])
    #ins.to_csv('test.csv')
    #print(ins['o'])
    ins['date'] =  ins['t']
    ins['value'] =  ins['v']
    ins = ins[['date','value']]
    data_list.append(ins)
result_data = data_list[0][['date']]
for i in range(len(data_list)):
    df = data_list[i]
    result_data = result_data.merge(df,how='left',on='date')
#last_data = result_data[(result_data.date>='2016-01-01') & (result_data.date<='2020-01-01')]
domain_data = result_data[(result_data.date>='2013-01-01')]
domain_data = domain_data.sort_values(by=['date'])
domain_data = domain_data.reset_index(drop=True)
domain_data = domain_data.merge(price_data,how='left',on=['date'])


f, axes = plt.subplots(figsize=(20, 10))
axes_fu = axes.twinx()
sns.lineplot(x="date", y="value",color='red',data=domain_data, ax=axes_fu)
sns.lineplot(x="date", y="close",color='black', data=domain_data, ax=axes)
axes.tick_params(labelsize=10)
plt.title('BTC Dominance', fontsize=10) 
axes.legend(loc='upper left', fontsize=5)
axes.set_xlabel('时间',fontsize=14,fontproperties=prop)
axes.set_ylabel("BTC price",fontsize=10)
axes_fu.set_ylabel("BTC Dominance",fontsize=10)

#plt.show()
plt.savefig('btc_dominance.png')
plt.close()
#add_mark(file = "btc_dominance.png", out = "out",mark = "MMC研究猿卡森出品", opacity=0.2, angle=30, space=30)


# ----- 矿工
url_address = ['https://api.glassnode.com/v1/metrics/transactions/transfers_volume_miners_to_exchanges']
url_name = ['k_fold']
# insert your API key here
API_KEY = '26BLocpWTcSU7sgqDdKzMHMpJDm'
data_list = []
for num in range(len(url_name)):
    print(num)
    addr = url_address[num]
    name = url_name[num]
    # make API request
    res_addr = requests.get(addr,params={'a': 'BTC', 'api_key': API_KEY})
    # convert to pandas dataframe
    ins = pd.read_json(res_addr.text, convert_dates=['t'])
    #ins.to_csv('test.csv')
    #print(ins['o'])
    ins['date'] =  ins['t']
    ins['value'] =  ins['v']
    ins = ins[['date','value']]
    data_list.append(ins)
result_data = data_list[0][['date']]
for i in range(len(data_list)):
    df = data_list[i]
    result_data = result_data.merge(df,how='left',on='date')
#last_data = result_data[(result_data.date>='2016-01-01') & (result_data.date<='2020-01-01')]
miner_data = result_data[(result_data.date>='2013-01-01')]
miner_data = miner_data.sort_values(by=['date'])
miner_data = miner_data.reset_index(drop=True)


date = []
miner_raw = []
for j in range(20,len(miner_data)+1):
    ins = miner_data[j-14:j]
    ins = ins.sort_values(by='date')
    ins = ins.reset_index(drop=True)
    date.append(ins['date'][13])
    miner_raw.append(np.mean(ins['value']))
miner_data_1 = pd.DataFrame({'date':date,'miner_raw':miner_raw})
miner_data_1 = miner_data_1[(miner_data_1.date>='2019-01-01')]


miner_data = miner_data_1.merge(price_data,how='left',on=['date'])


f, axes = plt.subplots(figsize=(20, 10))
axes_fu = axes.twinx()
sns.lineplot(x="date", y="miner_raw",color='red',data=miner_data, ax=axes_fu)
sns.lineplot(x="date", y="close",color='black', data=miner_data, ax=axes)
axes.tick_params(labelsize=10)
plt.title('Miner To Exhanges', fontsize=10) 
axes.legend(loc='upper left', fontsize=5)
axes.set_xlabel('时间',fontsize=14,fontproperties=prop)
axes.set_ylabel("BTC price",fontsize=10)
axes_fu.set_ylabel("BTC Volume",fontsize=10)

#plt.show()
plt.savefig('btc_miner_v.png')
plt.close()

#add_mark(file = "btc_miner_v.png", out = "out",mark = "MMC研究猿卡森出品", opacity=0.2, angle=30, space=30)



# ----- 巨鲸
url_address = [ 'https://api.glassnode.com/v1/metrics/transactions/transfers_volume_whales_to_exchanges_sum']
url_name = ['k_fold']
# insert your API key here
API_KEY = '26BLocpWTcSU7sgqDdKzMHMpJDm'
data_list = []
for num in range(len(url_name)):
    print(num)
    addr = url_address[num]
    name = url_name[num]
    # make API request
    res_addr = requests.get(addr,params={'a': 'BTC', 'api_key': API_KEY})
    # convert to pandas dataframe
    ins = pd.read_json(res_addr.text, convert_dates=['t'])
    #ins.to_csv('test.csv')
    #print(ins['o'])
    ins['date'] =  ins['t']
    ins['value'] =  ins['v']
    ins = ins[['date','value']]
    data_list.append(ins)
result_data = data_list[0][['date']]
for i in range(len(data_list)):
    df = data_list[i]
    result_data = result_data.merge(df,how='left',on='date')
#last_data = result_data[(result_data.date>='2016-01-01') & (result_data.date<='2020-01-01')]
whale_data = result_data[(result_data.date>='2013-01-01')]
whale_data = whale_data.sort_values(by=['date'])
whale_data = whale_data.reset_index(drop=True)
whale_data = whale_data.merge(price_data,how='left',on=['date'])


date = []
whale_raw = []
for j in range(30,len(whale_data)+1):
    ins = whale_data[j-30:j]
    ins = ins.sort_values(by='date')
    ins = ins.reset_index(drop=True)
    date.append(ins['date'][29])
    whale_raw.append(np.mean(ins['value']))
whale_data_1 = pd.DataFrame({'date':date,'whale_raw':whale_raw})
whale_data_1 = whale_data_1[(whale_data_1.date>='2019-01-01')]


whale_data = whale_data_1.merge(price_data,how='left',on=['date'])


f, axes = plt.subplots(figsize=(20, 10))
axes_fu = axes.twinx()
sns.lineplot(x="date", y="whale_raw",color='red',data=whale_data, ax=axes_fu)
sns.lineplot(x="date", y="close",color='black', data=whale_data, ax=axes)
axes.tick_params(labelsize=10)
plt.title('Whale To Exhanges', fontsize=10) 
axes.legend(loc='upper left', fontsize=5)
axes.set_xlabel('时间',fontsize=14,fontproperties=prop)
axes.set_ylabel("BTC price",fontsize=10)
axes_fu.set_ylabel("Whale Volume",fontsize=10)

#plt.show()
plt.savefig('btc_whale_v.png')
plt.close()
#add_mark(file = "btc_whale_v.png", out = "out",mark = "MMC研究猿卡森出品", opacity=0.2, angle=30, space=30)




# ---- btc净流入交易所

url_address = [ 'https://api.glassnode.com/v1/metrics/transactions/transfers_volume_exchanges_net']
url_name = ['k_fold']
# insert your API key here
API_KEY = '26BLocpWTcSU7sgqDdKzMHMpJDm'
data_list = []
for num in range(len(url_name)):
    print(num)
    addr = url_address[num]
    name = url_name[num]
    # make API request
    res_addr = requests.get(addr,params={'a': 'BTC', 'api_key': API_KEY})
    # convert to pandas dataframe
    ins = pd.read_json(res_addr.text, convert_dates=['t'])
    #ins.to_csv('test.csv')
    #print(ins['o'])
    ins['date'] =  ins['t']
    ins['value'] =  ins['v']
    ins = ins[['date','value']]
    data_list.append(ins)
result_data = data_list[0][['date']]
for i in range(len(data_list)):
    df = data_list[i]
    result_data = result_data.merge(df,how='left',on='date')
#last_data = result_data[(result_data.date>='2016-01-01') & (result_data.date<='2020-01-01')]
jinliuru_data = result_data[(result_data.date>='2013-01-01')]
jinliuru_data = jinliuru_data.sort_values(by=['date'])
jinliuru_data = jinliuru_data.reset_index(drop=True)
jinliuru_data = jinliuru_data.merge(price_data,how='left',on=['date'])
jinliuru_data = jinliuru_data[-120:]
f, axes = plt.subplots(figsize=(20, 10))
sns.barplot(x="date", y="value",data=jinliuru_data, ax=axes)
#sns.lineplot(x="date", y="close",color='black', data=jinliuru_data, ax=axes)
axes.tick_params(labelsize=10)
plt.title('BTC Netflow Exchanges', fontsize=10) 
axes.legend(loc='upper left', fontsize=5)
axes.set_xlabel('时间',fontsize=14,fontproperties=prop)
axes.set_ylabel("Netflow",fontsize=10)
#plt.show()
plt.savefig('btc_netflow.png')
plt.close()
#add_mark(file = "btc_netflow.png", out = "out",mark = "MMC研究猿卡森出品", opacity=0.2, angle=30, space=30)


# ----- usdt
url_address = [ 'https://api.glassnode.com/v1/metrics/transactions/transfers_volume_exchanges_net']
url_name = ['k_fold']
# insert your API key here
API_KEY = '26BLocpWTcSU7sgqDdKzMHMpJDm'
data_list = []
for num in range(len(url_name)):
    print(num)
    addr = url_address[num]
    name = url_name[num]
    # make API request
    res_addr = requests.get(addr,params={'a': 'USDT', 'api_key': API_KEY})
    # convert to pandas dataframe
    ins = pd.read_json(res_addr.text, convert_dates=['t'])
    #ins.to_csv('test.csv')
    #print(ins['o'])
    ins['date'] =  ins['t']
    ins['value'] =  ins['v']
    ins = ins[['date','value']]
    data_list.append(ins)
result_data = data_list[0][['date']]
for i in range(len(data_list)):
    df = data_list[i]
    result_data = result_data.merge(df,how='left',on='date')
#last_data = result_data[(result_data.date>='2016-01-01') & (result_data.date<='2020-01-01')]
u_jinliuru_data = result_data[(result_data.date>='2013-01-01')]
u_jinliuru_data = u_jinliuru_data.sort_values(by=['date'])
u_jinliuru_data = u_jinliuru_data.reset_index(drop=True)
u_jinliuru_data = u_jinliuru_data.merge(price_data,how='left',on=['date'])
u_jinliuru_data = u_jinliuru_data[-120:]
f, axes = plt.subplots(figsize=(20, 10))
sns.barplot(x="date", y="value",data=u_jinliuru_data, ax=axes)
#sns.lineplot(x="date", y="close",color='black', data=jinliuru_data, ax=axes)
axes.tick_params(labelsize=10)
plt.title('USDT Netflow Exchanges', fontsize=10) 
axes.legend(loc='upper left', fontsize=5)
axes.set_xlabel('时间',fontsize=14,fontproperties=prop)
axes.set_ylabel("Netflow",fontsize=10)
#plt.show()
plt.savefig('usdt_netflow.png')
plt.close()
#add_mark(file = "usdt_netflow.png", out = "out",mark = "MMC研究猿卡森出品", opacity=0.2, angle=30, space=30)

# btc_risk_index
url_address = [ 'https://api.glassnode.com/v1/metrics/signals/btc_risk_index']
url_name = ['k_fold']
# insert your API key here
API_KEY = '26BLocpWTcSU7sgqDdKzMHMpJDm'
data_list = []
for num in range(len(url_name)):
    print(num)
    addr = url_address[num]
    name = url_name[num]
    # make API request
    res_addr = requests.get(addr,params={'a': 'BTC', 'api_key': API_KEY})
    # convert to pandas dataframe
    ins = pd.read_json(res_addr.text, convert_dates=['t'])
    #ins.to_csv('test.csv')
    #print(ins['o'])
    ins['date'] =  ins['t']
    ins['value'] =  ins['v']
    ins = ins[['date','value']]
    data_list.append(ins)
result_data = data_list[0][['date']]
for i in range(len(data_list)):
    df = data_list[i]
    result_data = result_data.merge(df,how='left',on='date')
#last_data = result_data[(result_data.date>='2016-01-01') & (result_data.date<='2020-01-01')]
risk_data = result_data[(result_data.date>='2013-01-01')]
risk_data = risk_data.sort_values(by=['date'])
risk_data = risk_data.reset_index(drop=True)
risk_data = risk_data.merge(price_data,how='left',on=['date'])
risk_data = risk_data[risk_data.date>='2023-01-01']
risk_data['x1'] = 0.25
f, axes = plt.subplots(figsize=(20, 10))
axes_fu = axes.twinx()
sns.lineplot(x="date", y="value",data=risk_data, ax=axes)
sns.lineplot(x="date", y="x1",data=risk_data, ax=axes)
sns.lineplot(x="date", y="close",color='black', data=risk_data, ax=axes_fu)
axes.tick_params(labelsize=10)
plt.title('BTC RISK INDEX', fontsize=10) 
axes.legend(loc='upper left', fontsize=5)
axes.set_xlabel('时间',fontsize=14,fontproperties=prop)
axes.set_ylabel("INDEX",fontsize=10)
#plt.show()
plt.savefig('risk_index.png')
plt.close()
#add_mark(file = "risk_index.png", out = "out",mark = "MMC研究猿卡森出品", opacity=0.2, angle=30, space=30)


# ----- 基础btc数据
url_address = [ 'https://api.glassnode.com/v1/metrics/indicators/nupl_less_155_account_based']
url_name = ['k_fold']
# insert your API key here
API_KEY = '26BLocpWTcSU7sgqDdKzMHMpJDm'
data_list = []
for num in range(len(url_name)):
    print(num)
    addr = url_address[num]
    name = url_name[num]
    # make API request
    res_addr = requests.get(addr,params={'a': 'BTC', 'api_key': API_KEY})
    # convert to pandas dataframe
    ins = pd.read_json(res_addr.text, convert_dates=['t'])
    #ins.to_csv('test.csv')
    #print(ins['o'])
    ins['date'] =  ins['t']
    ins['value'] =  ins['v']
    ins = ins[['date','value']]
    data_list.append(ins)
result_data = data_list[0][['date']]
for i in range(len(data_list)):
    df = data_list[i]
    result_data = result_data.merge(df,how='left',on='date')
#last_data = result_data[(result_data.date>='2016-01-01') & (result_data.date<='2020-01-01')]
nupl_data = result_data[(result_data.date>='2013-01-01')]
nupl_data = nupl_data.sort_values(by=['date'])
nupl_data = nupl_data.reset_index(drop=True)
nupl_data = nupl_data.merge(price_data,how='left',on=['date'])
nupl_data = nupl_data[nupl_data.date>='2023-01-01']
nupl_data['x1'] = 0
nupl_data['x2'] = 0.25
f, axes = plt.subplots(figsize=(20, 10))
axes_fu = axes.twinx()
sns.lineplot(x="date", y="value",color='red',data=nupl_data, ax=axes)
sns.lineplot(x="date", y="x1",color='green',linestyle='--',data=nupl_data, ax=axes)
sns.lineplot(x="date", y="x2",color='green',linestyle='--',data=nupl_data, ax=axes)
sns.lineplot(x="date", y="close",color='black', data=nupl_data, ax=axes_fu)
axes.tick_params(labelsize=10)
plt.title('Entity-Adjusted STH-NUPL', fontsize=10) 
axes.legend(loc='upper left', fontsize=5)
axes.set_xlabel('时间',fontsize=14,fontproperties=prop)
axes.set_ylabel("STH-NUPL",fontsize=10)
#plt.show()
plt.savefig('sth_nupl.png')
plt.close()
#add_mark(file = "sth_nupl.png", out = "out",mark = "MMC研究猿卡森出品", opacity=0.2, angle=30, space=30)

# =======================================================================黑天鹅地址监控===================================================================

url_address = [ 'https://api.glassnode.com/v1/metrics/distribution/balance_us_government','https://api.glassnode.com/v1/metrics/distribution/balance_mtgox_trustee','https://api.glassnode.com/v1/metrics/market/price_usd_close']
url_name = ['balance_us','balance_mtgox','Price']
# insert your API key here
API_KEY = '26BLocpWTcSU7sgqDdKzMHMpJDm'
data_list = []
for num in range(len(url_name)):
    print(num)
    addr = url_address[num]
    name = url_name[num]
    # make API request
    res_addr = requests.get(addr,params={'a': 'BTC', 'api_key': API_KEY})
    # convert to pandas dataframe
    ins = pd.read_json(res_addr.text, convert_dates=['t'])
    #ins.to_csv('test.csv')
    #print(ins['o'])
    ins['date'] =  ins['t']
    ins['value'] =  ins['v']
    ins = ins[['date','value']]
    data_list.append(ins)
result_data = data_list[0][['date']]
for i in range(len(data_list)):
    df = data_list[i]
    result_data = result_data.merge(df,how='left',on='date')
#last_data = result_data[(result_data.date>='2016-01-01') & (result_data.date<='2020-01-01')]
last_data = result_data[(result_data.date>='2018-01-01')]
last_data['log(BTC price)'] = np.log(last_data['value'])
last_data.rename(columns={'value_x':'Balance US','value_y':'Balance Mt.GOX'},inplace=True)
last_data = last_data.sort_values(by=['date'])
last_data = last_data.reset_index(drop=True)

last_data_us = last_data[['date','Balance US']]
last_data_us['nb'] = last_data_us['Balance US'].shift(-1)
last_data_us['nd'] = last_data_us['date'].shift(-1)
last_data_us = last_data_us[0:-1]
last_data_us['cha'] = last_data_us['nb'] - last_data_us['Balance US']
last_data_us = last_data_us[last_data_us.cha < 0]
last_data_us = last_data_us[['nd','nb','Balance US']]

us_pre_b = last_data['Balance US'][len(last_data)-2]
us_pre_t = last_data['Balance US'][len(last_data)-1]

gox_pre_b = last_data['Balance Mt.GOX'][len(last_data)-2]
gox_pre_t = last_data['Balance Mt.GOX'][len(last_data)-1]


change_us = us_pre_t - us_pre_b
if change_us == 0:
    content_us = '美国政府持有地址目前有%s个BTC，相比昨日无变化.'%(str(us_pre_t))
elif change_us < 0:
    content_us = '美国政府持有地址目前有%s个BTC，相比昨日减少%s个BTC.'%(str(us_pre_t),str(-change_us))
else:
    content_us = '美国政府持有地址目前有%s个BTC，相比昨日增加%s个BTC.'%(str(us_pre_t),str(change_us))
print(content_us)
change_gox = gox_pre_t - gox_pre_b
if change_gox == 0:
    content_gox = 'MT.GOX持有地址目前有%s个BTC，相比昨日无变化.'%(str(gox_pre_t))
elif change_gox < 0:
    content_gox = 'MT.GOX持有地址目前有%s个BTC，相比昨日减少%s个BTC.'%(str(gox_pre_t),str(-change_gox))
else:
    content_gox = 'MT.GOX持有地址目前有%s个BTC，相比昨日增加%s个BTC.'%(str(gox_pre_t),str(change_gox))
f, axes = plt.subplots(figsize=(20, 10))
axes_fu = axes.twinx()

sns.lineplot(x="date", y="log(BTC price)", data=last_data, color = 'black', linewidth=0.5,ci=95, ax=axes)
sns.lineplot(x="date", y="Balance US",color='red', linewidth=0.5,data=last_data, ci=95, ax=axes_fu)
sns.lineplot(x="date", y="Balance Mt.GOX",color='green', linewidth=0.5,data=last_data, ci=95, ax=axes_fu)
axes.set_title('US GOV和MT GOX持币地址BTC数量', fontsize=30,fontproperties=prop) 
axes.legend(loc='upper left', fontsize=10)
axes.set_xlabel('时间',fontsize=14,fontproperties=prop)
axes.set_ylabel("BTC价格",fontsize=14,fontproperties=prop)
axes_fu.set_ylabel("持有BTC总量",fontsize=14,fontproperties=prop)
plt.savefig('US_MT.png',  bbox_inches='tight')
plt.close()

#add_mark(file = "US_MT.png", out = "out",mark = "MMC研究猿卡森出品", opacity=0.2, angle=30, space=30)


# =======================================================================链上的4种稳定币===================================================================

url_address = [ 'https://api.glassnode.com/v1/metrics/supply/current']
url_name = ['supply_usdt']
# insert your API key here
API_KEY = '26BLocpWTcSU7sgqDdKzMHMpJDm'
data_list = []
for num in range(len(url_name)):
    print(num)
    addr = url_address[num]
    name = url_name[num]
    # make API request
    res_addr = requests.get(addr,params={'a': 'USDT', 'api_key': API_KEY})
    # convert to pandas dataframe
    ins = pd.read_json(res_addr.text, convert_dates=['t'])
    #ins.to_csv('test.csv')
    #print(ins['o'])
    ins['date'] =  ins['t']
    ins['value'] =  ins['v']
    ins = ins[['date','value']]
    data_list.append(ins)
result_data = data_list[0][['date']]
for i in range(len(data_list)):
    df = data_list[i]
    result_data = result_data.merge(df,how='left',on='date')
#last_data = result_data[(result_data.date>='2016-01-01') & (result_data.date<='2020-01-01')]
usdt_data = result_data[(result_data.date>='2018-01-01')]
usdt_data.rename(columns={'value':'supply_usdt'},inplace=True)
'''
url_address = [ 'https://api.glassnode.com/v1/metrics/supply/current']
url_name = ['supply_usdc']
# insert your API key here
API_KEY = '26BLocpWTcSU7sgqDdKzMHMpJDm'
data_list = []
for num in range(len(url_name)):
    print(num)
    addr = url_address[num]
    name = url_name[num]
    # make API request
    res_addr = requests.get(addr,params={'a': 'USDC', 'api_key': API_KEY})
    # convert to pandas dataframe
    ins = pd.read_json(res_addr.text, convert_dates=['t'])
    #ins.to_csv('test.csv')
    #print(ins['o'])
    ins['date'] =  ins['t']
    ins['value'] =  ins['v']
    ins = ins[['date','value']]
    data_list.append(ins)
result_data = data_list[0][['date']]
for i in range(len(data_list)):
    df = data_list[i]
    result_data = result_data.merge(df,how='left',on='date')
#last_data = result_data[(result_data.date>='2016-01-01') & (result_data.date<='2020-01-01')]
usdc_data = result_data[(result_data.date>='2018-01-01')]
usdc_data.rename(columns={'value':'supply_usdc'},inplace=True)
url_address = [ 'https://api.glassnode.com/v1/metrics/supply/current']
url_name = ['supply_busd']
# insert your API key here
API_KEY = '26BLocpWTcSU7sgqDdKzMHMpJDm'
data_list = []
for num in range(len(url_name)):
    print(num)
    addr = url_address[num]
    name = url_name[num]
    # make API request
    res_addr = requests.get(addr,params={'a': 'BUSD', 'api_key': API_KEY})
    # convert to pandas dataframe
    ins = pd.read_json(res_addr.text, convert_dates=['t'])
    #ins.to_csv('test.csv')
    #print(ins['o'])
    ins['date'] =  ins['t']
    ins['value'] =  ins['v']
    ins = ins[['date','value']]
    data_list.append(ins)
result_data = data_list[0][['date']]
for i in range(len(data_list)):
    df = data_list[i]
    result_data = result_data.merge(df,how='left',on='date')
#last_data = result_data[(result_data.date>='2016-01-01') & (result_data.date<='2020-01-01')]
busd_data = result_data[(result_data.date>='2018-01-01')]
busd_data.rename(columns={'value':'supply_busd'},inplace=True)
'''
url_address = [ 'https://api.glassnode.com/v1/metrics/supply/current']
url_name = ['supply_tusd']
# insert your API key here
API_KEY = '26BLocpWTcSU7sgqDdKzMHMpJDm'
data_list = []
for num in range(len(url_name)):
    print(num)
    addr = url_address[num]
    name = url_name[num]
    # make API request
    res_addr = requests.get(addr,params={'a': 'TUSD', 'api_key': API_KEY})
    # convert to pandas dataframe
    ins = pd.read_json(res_addr.text, convert_dates=['t'])
    #ins.to_csv('test.csv')
    #print(ins['o'])
    ins['date'] =  ins['t']
    ins['value'] =  ins['v']
    ins = ins[['date','value']]
    data_list.append(ins)
result_data = data_list[0][['date']]
for i in range(len(data_list)):
    df = data_list[i]
    result_data = result_data.merge(df,how='left',on='date')
#last_data = result_data[(result_data.date>='2016-01-01') & (result_data.date<='2020-01-01')]
tusd_data = result_data[(result_data.date>='2018-01-01')]
tusd_data.rename(columns={'value':'supply_tusd'},inplace=True)
#all_stable_coin = usdt_data.merge(usdc_data,how='left',on=['date']).merge(busd_data,how='left',on=['date']).merge(tusd_data,how='left',on=['date'])
all_stable_coin = usdt_data.merge(tusd_data,how='left',on=['date'])
all_stable_coin = all_stable_coin.fillna(0)
#all_stable_coin['all'] = all_stable_coin['supply_usdt'] + all_stable_coin['supply_usdc'] + all_stable_coin['supply_busd'] + all_stable_coin['supply_tusd']
all_stable_coin['all'] = all_stable_coin['supply_usdt'] + all_stable_coin['supply_tusd']
all_stable_coin = all_stable_coin.merge(price_data,how='left',on=['date'])
all_stable_coin = all_stable_coin.sort_values(by='date')
all_stable_coin = all_stable_coin.reset_index(drop=True)
all_stable_coin['Toal Supply'] = all_stable_coin['all']/100000000
all_stable_coin = all_stable_coin[all_stable_coin.date>='2022-01-01']
f, axes = plt.subplots(figsize=(20, 10))
axes_fu = axes.twinx()

sns.lineplot(x="date", y="close",color='black', linewidth=0.5,data=all_stable_coin, ci=95, ax=axes)
sns.lineplot(x="date", y="Toal Supply",color='blue', linewidth=0.5,data=all_stable_coin, ci=95, ax=axes_fu)
#plt.title('USDT/USDC/BUSD/TUSD链上总发行量', fontsize=20,fontproperties=prop) 
plt.title('USDT/TUSD链上总发行量', fontsize=20,fontproperties=prop) 
axes.set_xlabel('时间',fontsize=14,fontproperties=prop)
axes.set_ylabel("BTC价格",fontsize=14,fontproperties=prop)
axes_fu.set_ylabel("发行总量",fontsize=14,fontproperties=prop)
plt.savefig('BTC价格-稳定币发行总量.png',  bbox_inches='tight')
plt.close()

#add_mark(file = "BTC价格-稳定币发行总量.png", out = "out",mark = "MMC研究猿卡森出品", opacity=0.2, angle=30, space=30)




# =======================================================================链上指标===================================================================

def cal(x):
    if x>= pd.to_datetime('2013-01-01') and x<= pd.to_datetime('2016-10-31'):
        y = 'Second cycle'
    elif x>= pd.to_datetime('2016-11-01') and x<= pd.to_datetime('2020-04-30'):
        y = 'Third cycle'
    else:
        y = 'Fourth cycle'
    return y
url_address = ['https://api.glassnode.com/v1/metrics/indicators/puell_multiple',
                'https://api.glassnode.com/v1/metrics/indicators/sopr_adjusted',
                'https://api.glassnode.com/v1/metrics/market/mvrv_z_score',
                'https://api.glassnode.com/v1/metrics/indicators/rhodl_ratio',
                'https://api.glassnode.com/v1/metrics/indicators/net_realized_profit_loss',
                'https://api.glassnode.com/v1/metrics/market/price_usd_close',
                'https://api.glassnode.com/v1/metrics/supply/profit_relative',
                'https://api.glassnode.com/v1/metrics/transactions/transfers_volume_to_exchanges_sum',
                'https://api.glassnode.com/v1/metrics/transactions/transfers_volume_from_exchanges_sum']
url_name = ['Puell Multiple', 'aSOPR','MVRV Z-Score','RHODL Ratio','Net Realized Profit/Loss','Price','Percent Supply in Profit','in_exchanges', 'out_exchanges']
# insert your API key here
API_KEY = '26BLocpWTcSU7sgqDdKzMHMpJDm'
data_list = []
for num in range(len(url_name)):
    print(num)
    addr = url_address[num]
    name = url_name[num]
    # make API request
    res_addr = requests.get(addr,params={'a': 'BTC', 'api_key': API_KEY})
    # convert to pandas dataframe
    ins = pd.read_json(res_addr.text, convert_dates=['t'])
    ins['date'] =  ins['t']
    ins[name] =  ins['v']
    ins = ins[['date',name]]
    data_list.append(ins)

result_data = data_list[0][['date']]
for i in range(len(data_list)):
    df = data_list[i]
    result_data = result_data.merge(df,how='left',on='date')
#last_data = result_data[(result_data.date>='2016-01-01') & (result_data.date<='2020-01-01')]
last_data = result_data[(result_data.date>='2012-10-01')]
last_data = last_data.sort_values(by=['date'])
last_data = last_data.reset_index(drop=True)
date = []
pm = []
mvrv = []
rhold = []
price = []
sopr_7 = []
sopr = []
supply = []
for j in range(len(last_data)-49):
    ins = last_data[j:j+50]
    ins = ins.reset_index(drop=True)
    date.append(ins['date'][49])
    pm.append(ins['Puell Multiple'][49])
    mvrv.append(ins['MVRV Z-Score'][49])
    rhold.append(ins['RHODL Ratio'][49])
    sopr.append(ins['aSOPR'][49])
    supply.append(ins['Percent Supply in Profit'][49])
    price.append(ins['Price'][49])
    #短期指标
    sopr_7.append(np.mean(ins['aSOPR'][-7:]))
res_df = pd.DataFrame({'date':date,'Puell Multiple':pm,'MVRV Z-Score':mvrv,'RHODL Ratio':rhold,'Price':price,'Percent Supply in Profit':supply,'7MA aSOPR':sopr_7,'aSOPR':sopr})
res_df = res_df[(res_df.date>='2019-01-01')]
res_df['cycle'] = res_df['date'].apply(lambda x:cal(x))
res_df['log(BTC price)'] = np.log(res_df['Price'])
res_df['log(RHODL Ratio)'] = np.log(res_df['RHODL Ratio'])
res_df['x1'] = 7
res_df['x2'] = 0
res_df['y1'] = 4
res_df['y2'] = 0.5
res_df['z1'] = np.log(49000)
res_df['z2'] = np.log(350)
res_df['w'] = 1
res_df['p1'] = 0.9
res_df['p2'] = 0.5

res_df['f1'] = 0.95
res_df['f2'] = 0.5

# --- asopr
f, axes = plt.subplots(figsize=(20, 10))
axes_fu = axes.twinx()
sns.lineplot(x="date", y="w", data=res_df, color='red',linestyle='--', ax=axes_fu)
sns.lineplot(x="date", y="aSOPR",color='black',data=res_df, ax=axes_fu)
sns.lineplot(x="date", y="log(BTC price)",hue = 'cycle', data=res_df, ax=axes)
axes.tick_params(labelsize=10)
axes.legend(loc='upper left', fontsize=5)


axes.set_xlabel('时间',fontsize=14,fontproperties=prop)
axes.set_ylabel("log(BTC price)",fontsize=10)
axes_fu.set_ylabel("aSOPR",fontsize=10)
plt.title('每日aSOPR值', fontsize=20,fontproperties=prop) 

plt.savefig('aSOPR.png') 
plt.close()

#add_mark(file = "aSOPR.png", out = "out",mark = "MMC研究猿卡森出品", opacity=0.2, angle=30, space=30)

# --- 7ma asopr
f, axes = plt.subplots(figsize=(20, 10))
axes_fu = axes.twinx()
sns.lineplot(x="date", y="w", data=res_df, color='red',linestyle='--', ax=axes_fu)
sns.lineplot(x="date", y="7MA aSOPR",color='black',data=res_df, ax=axes_fu)
sns.lineplot(x="date", y="log(BTC price)",hue = 'cycle', data=res_df, ax=axes)
axes.tick_params(labelsize=10)
axes.legend(loc='upper left', fontsize=5)
plt.title('7MA aSOPR值', fontsize=10) 
axes.set_xlabel('时间',fontsize=14,fontproperties=prop)
axes.set_ylabel("log(BTC price)",fontsize=10)
axes_fu.set_ylabel("7MA aSOPR",fontsize=10)

plt.savefig('7MA_aSOPR.png') 
plt.close()

#add_mark(file = "7MA_aSOPR.png", out = "out",mark = "MMC研究猿卡森出品", opacity=0.2, angle=30, space=30)

# --- mvrv
f, axes = plt.subplots(figsize=(20, 10))
axes_fu = axes.twinx()
sns.lineplot(x="date", y="x1", data=res_df, color = 'green',linestyle='--',ax=axes_fu)
sns.lineplot(x="date", y="x2", data=res_df, color='red',linestyle='--',ax=axes_fu)
sns.lineplot(x="date", y="MVRV Z-Score",color='black',data=res_df,ax=axes_fu)
sns.lineplot(x="date", y="log(BTC price)",hue = 'cycle', data=res_df,ax=axes)
axes.tick_params(labelsize=10)
plt.title('MVRV Z-Score值', fontsize=10) 
axes.legend(loc='upper left', fontsize=5)
axes.set_xlabel('时间',fontsize=14,fontproperties=prop)
axes.set_ylabel("log(BTC price)",fontsize=10)
axes_fu.set_ylabel("MVRV Z-Score",fontsize=10)

plt.savefig('MVRV_Z_Score.png') 
plt.close()
#add_mark(file = "MVRV_Z_Score.png", out = "out",mark = "MMC研究猿卡森出品", opacity=0.2, angle=30, space=30)

# --- pm
f, axes = plt.subplots(figsize=(20, 10))
axes_fu = axes.twinx()
sns.lineplot(x="date", y="y1", data=res_df, color = 'green',linestyle='--', ax=axes_fu)
sns.lineplot(x="date", y="y2", data=res_df, color='red', linestyle='--',ax=axes_fu)
sns.lineplot(x="date", y="Puell Multiple",color='black',data=res_df, ax=axes_fu)
sns.lineplot(x="date", y="log(BTC price)",hue = 'cycle', data=res_df, ax=axes)
axes.tick_params(labelsize=10)
plt.title('Puell Multiple值', fontsize=10) 
axes.legend(loc='upper left', fontsize=5)

axes.set_xlabel('时间',fontsize=14,fontproperties=prop)
axes.set_ylabel("log(BTC price)",fontsize=10)
axes_fu.set_ylabel("Puell Multiple",fontsize=10)

#plt.show()
plt.savefig('Puell.png')
plt.close()
#add_mark(file = "Puell.png", out = "out",mark = "MMC研究猿卡森出品", opacity=0.2, angle=30, space=30)

# --- rhold
f, axes = plt.subplots(figsize=(20, 10))
axes_fu = axes.twinx()
sns.lineplot(x="date", y="z1", data=res_df, color = 'green',linestyle='--',ax=axes_fu)
sns.lineplot(x="date", y="z2", data=res_df, color='red', linestyle='--',ax=axes_fu)
sns.lineplot(x="date", y="log(RHODL Ratio)",color='black',data=res_df, ax=axes_fu)
sns.lineplot(x="date", y="log(BTC price)",hue = 'cycle', data=res_df, ax=axes)
axes.tick_params(labelsize=10)
plt.title('log(RHODL Ratio)值', fontsize=10) 
axes.legend(loc='upper left', fontsize=5)

axes.set_xlabel('时间',fontsize=14,fontproperties=prop)
axes.set_ylabel("log(BTC price)",fontsize=10)
axes_fu.set_ylabel("log(RHODL Ratio)",fontsize=10)

#plt.show()
plt.savefig('RHODL.png')
plt.close()
#add_mark(file = "RHODL.png", out = "out",mark = "MMC研究猿卡森出品", opacity=0.2, angle=30, space=30)

# --- supply

f, axes = plt.subplots(figsize=(20, 10))
axes_fu = axes.twinx()
sns.lineplot(x="date", y="f2", data=res_df, color='red',linestyle='--', ax=axes_fu)
sns.lineplot(x="date", y="f1", data=res_df, color='green',linestyle='--', ax=axes_fu)
sns.lineplot(x="date", y="Percent Supply in Profit",color='black',data=res_df, ax=axes_fu)
sns.lineplot(x="date", y="log(BTC price)",hue = 'cycle', data=res_df, ax=axes)
axes.tick_params(labelsize=10)
plt.title('Percent Supply in Profit值', fontsize=10) 
axes.legend(loc='upper left', fontsize=5)
axes.set_xlabel('时间',fontsize=14,fontproperties=prop)
axes.set_ylabel("log(BTC price)",fontsize=10)
axes_fu.set_ylabel("Percent Supply in Profit",fontsize=10)

#plt.show()
plt.savefig('Percent_Supply.png')
plt.close()
#add_mark(file = "Percent_Supply.png", out = "out",mark = "MMC研究猿卡森出品", opacity=0.2, angle=30, space=30)
# =======================================================================未来24小时价格预测==================================================================

date_now = datetime.datetime.utcnow()
date_before = pd.to_datetime(date_now) - datetime.timedelta(days=8)
# 读取原始数据
raw_data = pd.read_csv('/root/usdt/usdt.csv')
raw_data = raw_data[2:]
raw_data['date'] = raw_data['date'].apply(lambda x: str(x)[6:10] + '/' + str(x)[3:5] + '/' + str(x)[0:2] + ' ' + str(x)[11:19])
raw_data['date'] = pd.to_datetime(raw_data['date'])
raw_data['date'] = raw_data['date'] - datetime.timedelta(hours=8)
#取最近7天的数据
raw_data = raw_data[(raw_data.date >= date_before) & (raw_data.date < date_now)]
raw_data['date_dd'] = raw_data['date'].apply(lambda x: str(x)[0:10])
raw_data['date_hh'] = raw_data['date'].apply(lambda x: str(x)[11:13])
raw_data['per'] = raw_data['usdt']/raw_data['usd']
usdt_data = raw_data
# =======================================================================衍生品指标=====================================================================
url_address = ['https://api.glassnode.com/v1/metrics/derivatives/options_25delta_skew_1_week',
                'https://api.glassnode.com/v1/metrics/derivatives/futures_volume_daily_perpetual_sum']
url_name = ['options', 'futures']
# insert your API key here
API_KEY = '26BLocpWTcSU7sgqDdKzMHMpJDm'
data_list = []
for num in range(len(url_name)):
    print(num)
    addr = url_address[num]
    name = url_name[num]
    # make API request
    res_addr = requests.get(addr,params={'a': 'BTC', 'api_key': API_KEY})
    # convert to pandas dataframe
    ins = pd.read_json(res_addr.text, convert_dates=['t'])
    ins['date'] =  ins['t']
    ins[name] =  ins['v']
    ins = ins[['date',name]]
    data_list.append(ins)

result_data = data_list[0][['date']]
for i in range(len(data_list)):
    df = data_list[i]
    result_data = result_data.merge(df,how='left',on='date')
#last_data = result_data[(result_data.date>='2016-01-01') & (result_data.date<='2020-01-01')]
dev_data = result_data[(result_data.date>='2020-01-01')]
dev_data = dev_data.sort_values(by=['date'])
dev_data = dev_data.reset_index(drop=True)



date = []
dev_op = []
dev_fu = []
for j in range(7,len(dev_data)+1):
    ins = dev_data[j-7:j]
    ins = ins.sort_values(by='date')
    ins = ins.reset_index(drop=True)
    date.append(ins['date'][6])
    dev_op.append(np.mean(ins['options']))
    dev_fu.append(np.mean(ins['futures']))
dev_data_1 = pd.DataFrame({'date':date,'dev_op':dev_op,'dev_fu':dev_fu})
dev_data_1 = dev_data_1[(dev_data_1.date>='2021-01-01')]


dev_data_1 = dev_data_1.merge(price_data,how='left',on=['date'])


f, axes = plt.subplots(figsize=(20, 10))
axes_fu = axes.twinx()
sns.lineplot(x="date", y="dev_op",color='red',data=dev_data_1, ax=axes_fu)
sns.lineplot(x="date", y="close",color='black', data=dev_data_1, ax=axes)
axes.tick_params(labelsize=10)
plt.title('Options 25 Delta Skew (1 Week) - Deribit (7d Moving Average)', fontsize=10) 
axes.legend(loc='upper left', fontsize=5)
axes.set_xlabel('时间',fontsize=14,fontproperties=prop)
axes.set_ylabel("BTC price",fontsize=10)
axes_fu.set_ylabel("Options 25 Delta Skew (1 Week)",fontsize=10)

#plt.show()
plt.savefig('btc_options.png')
plt.close()


f, axes = plt.subplots(figsize=(20, 10))
axes_fu = axes.twinx()
sns.lineplot(x="date", y="dev_fu",color='red',data=dev_data_1, ax=axes_fu)
sns.lineplot(x="date", y="close",color='black', data=dev_data_1, ax=axes)
axes.tick_params(labelsize=10)
plt.title('Futures Volume Perpetual [USD] - All Exchanges (7d Moving Average)', fontsize=10) 
axes.legend(loc='upper left', fontsize=5)
axes.set_xlabel('时间',fontsize=14,fontproperties=prop)
axes.set_ylabel("BTC price",fontsize=10)
axes_fu.set_ylabel("Futures Volume Perpetual",fontsize=10)

#plt.show()
plt.savefig('btc_futures.png')
plt.close()






# ========================================================================附录=========================================================================
url_address = ['https://api.glassnode.com/v1/metrics/market/price_usd_close']
url_name = ['Price']
# insert your API key here
API_KEY = '26BLocpWTcSU7sgqDdKzMHMpJDm'
data_list = []
for num in range(len(url_name)):
    print(num)
    addr = url_address[num]
    name = url_name[num]
    # make API request
    res_addr = requests.get(addr,params={'a': 'BTC', 'api_key': API_KEY})
    # convert to pandas dataframe
    ins = pd.read_json(res_addr.text, convert_dates=['t'])
    ins['date'] =  ins['t']
    ins[name] =  ins['v']
    ins = ins[['date',name]]
    data_list.append(ins)

result_data = data_list[0][['date']]
for i in range(len(data_list)):
    df = data_list[i]
    result_data = result_data.merge(df,how='left',on='date')
#last_data = result_data[(result_data.date>='2016-01-01') & (result_data.date<='2020-01-01')]
btc_data = result_data[(result_data.date>='2019-10-01')]
btc_data = btc_data.sort_values(by=['date'])
btc_data = btc_data.reset_index(drop=True)
data_list = []
for num in range(len(url_name)):
    print(num)
    addr = url_address[num]
    name = url_name[num]
    # make API request
    res_addr = requests.get(addr,params={'a': 'ETH', 'api_key': API_KEY})
    # convert to pandas dataframe
    ins = pd.read_json(res_addr.text, convert_dates=['t'])
    ins['date'] =  ins['t']
    ins[name] =  ins['v']
    ins = ins[['date',name]]
    data_list.append(ins)

result_data = data_list[0][['date']]
for i in range(len(data_list)):
    df = data_list[i]
    result_data = result_data.merge(df,how='left',on='date')
#last_data = result_data[(result_data.date>='2016-01-01') & (result_data.date<='2020-01-01')]
eth_data = result_data[(result_data.date>='2019-01-01')]
eth_data = eth_data.sort_values(by=['date'])
eth_data = eth_data.reset_index(drop=True)
combine_data = eth_data.merge(btc_data,how='left',on=['date'])
combine_data['per'] = combine_data['Price_x']/combine_data['Price_y']


f, axes = plt.subplots(figsize=(20, 10))
axes_fu = axes.twinx()
sns.lineplot(x="date", y="per",color='red',data=combine_data, ax=axes_fu)
sns.lineplot(x="date", y="Price_y",color='black', data=combine_data, ax=axes)
axes.tick_params(labelsize=10)
plt.title('ETH Price/BTC Price', fontsize=10) 
axes.legend(loc='upper left', fontsize=5)
axes.set_xlabel('时间',fontsize=14,fontproperties=prop)
axes.set_ylabel("BTC price",fontsize=10)
axes_fu.set_ylabel("ETH Price/BTC Price",fontsize=10)

#plt.show()
plt.savefig('eth_btc.png')
plt.close()
#add_mark(file = "eth_btc.png", out = "out",mark = "MMC研究猿卡森出品", opacity=0.2, angle=30, space=30)

# =======================================================================输出docment===================================================================
import sys
from docx import Document
from docx.shared import Inches
# 创建文档对象
document = Document()
# 设置文档标题，中文要用unicode字符串
document.add_heading(u'%s日BTC链上数据一览'%(str(date_now)[0:10]),0)
# 添加有序列表

document.add_paragraph('BTC价格和流入流出交易所',style = 'ListBullet')
document.add_paragraph('链上稳定币监控',style = 'ListBullet')
document.add_paragraph('重要链上指标监控',style = 'ListBullet')
document.add_paragraph('黑天鹅事件监控',style = 'ListBullet')
document.add_paragraph('衍生品指标监控',style = 'ListBullet')

# -----  120/200/4y 均线
document.add_page_break()
#p = document.add_paragraph('This is a paragraph in new page.')
document.add_heading(u'1.BTC价格和流入流出交易所',level = 1)
jun_df = jun_df.reset_index(drop=True)
value_1 = jun_df['price_raw'][len(jun_df)-1]
value_120 = jun_df['price_ma120'][len(jun_df)-1]
value_200 = jun_df['price_ma200'][len(jun_df)-1]
value_4y = jun_df['price_ma4y'][len(jun_df)-1]

btc_zhangdiefu = round( ((price_data['close'][len(price_data)-2] - price_data['close'][len(price_data)-3])/price_data['close'][len(price_data)-3])*100,2)
btc_zhangdiefu_1 = str(btc_zhangdiefu)+'%'

document.add_paragraph('昨日BTC收盘价格为：%s'%(str(value_1)),style = 'ListBullet')
document.add_paragraph('昨日BTC涨幅为：%s'%(btc_zhangdiefu_1),style = 'ListBullet')
document.add_paragraph('昨日BTC收盘MA120价格为：%s'%(str(value_120)),style = 'ListBullet')
document.add_paragraph('昨日BTC收盘MA200价格为：%s'%(str(value_200)),style = 'ListBullet')
document.add_paragraph('昨日BTC收盘MA4Y价格为：%s'%(str(value_4y)),style = 'ListBullet')
# 添加图片，并指定宽度
document.add_picture('/root/xianhuo/BTC价格120D.png',width = Inches(5.25))

document.add_paragraph('BTC市值占比',style = 'ListBullet')
document.add_paragraph('牛市要启动，比特币市值占比至少要高于50%。',style = 'ListBullet')
btc_dominance = str(domain_data['value'][len(domain_data)-1]) + '%'
text = '今日BTC市值占比为：%s。'%(btc_dominance)
document.add_paragraph(text,style = 'ListBullet')
# 添加图片，并指定宽度
document.add_picture('/root/xianhuo/btc_dominance.png',width = Inches(5.25))

document.add_paragraph('交易所比特币净流入量',style = 'ListBullet')
document.add_paragraph('交易所比特币净流入量为负，说明有人在提币，净流入量为正，说明有人在卖币，在价格低位的时候一定是有人长期在定投囤币。',style = 'ListBullet')
jinliuru_data = jinliuru_data.reset_index(drop=True)

btc_netflow = jinliuru_data['value'][len(jinliuru_data)-1]
text = '昨日交易所比特币净流入量为：%s。'%(btc_netflow)
document.add_paragraph(text,style = 'ListBullet')
# 添加图片，并指定宽度
document.add_picture('/root/xianhuo/btc_netflow.png',width = Inches(5.25))

document.add_paragraph('近14日矿工每日平均流入交易所比特币量',style = 'ListBullet')
document.add_paragraph('矿工转入交易所的BTC数量在减少说明长期看涨，在增多说明长期看跌。',style = 'ListBullet')
miner_data = miner_data.reset_index(drop=True)

miner_netflow = miner_data['miner_raw'][len(miner_data)-1]
text = '昨日矿工流入交易所比特币量为：%s。'%(miner_netflow)
document.add_paragraph(text,style = 'ListBullet')
# 添加图片，并指定宽度
document.add_picture('/root/xianhuo/btc_miner_v.png',width = Inches(5.25))

document.add_paragraph('近30日巨鲸每日平均流入交易所比特币量',style = 'ListBullet')
document.add_paragraph('巨鲸转入交易所的BTC数量在减少说明长期看涨，在增多说明长期看跌。',style = 'ListBullet')
whale_data = whale_data.reset_index(drop=True)

whale_netflow = whale_data['whale_raw'][len(whale_data)-1]
text = '昨日巨鲸流入交易所比特币量为：%s。'%(whale_netflow)
document.add_paragraph(text,style = 'ListBullet')
# 添加图片，并指定宽度
document.add_picture('/root/xianhuo/btc_whale_v.png',width = Inches(5.25))


# -----  链上稳定币监控
document.add_page_break()
#p = document.add_paragraph('This is a paragraph in new page.')
document.add_heading(u'2.链上稳定币监控',level = 1)

document.add_paragraph('以下统计的稳定币包括USDT/TUSD两种',style = 'ListBullet')
document.add_picture('/root/xianhuo/BTC价格-稳定币发行总量.png',width = Inches(6))

document.add_paragraph('近7日稳定币总量值',style = 'ListBullet')
all_stable_coin = all_stable_coin[['date','Toal Supply']]
all_stable_coin = all_stable_coin.reset_index(drop=True)
sub_all_stable_coin = all_stable_coin[-7:]

t = document.add_table(rows=1, cols=2) # 插入表格，先将表头写好，参数：rows:行，cols:列
hdr_cells = t.rows[0].cells
hdr_cells[0].text = '时间' # 表头
hdr_cells[1].text = '数量'# 表头

for d in sub_all_stable_coin.values.tolist(): # 
    print(d)
    for date,Supply in [d]: # 读取每一行内容
        row_cells = t.add_row().cells # 读到一行就在word的表格中插入一行
        row_cells[0].text = str(date)
        row_cells[1].text = str(round(Supply,2))
        
#usdt的溢价率
document.add_paragraph('USDT场外溢价率监控',style = 'ListBullet')
#p = document.add_paragraph('This is a paragraph in new page.')

usdt_data = usdt_data.sort_values(by='date')
usdt_data = usdt_data.reset_index(drop=True)

usdt_price = usdt_data['usdt'][len(usdt_data)-1]
usd_price = usdt_data['usd'][len(usdt_data)-1]
usdt_per = usdt_data['per'][len(usdt_data)-1]
document.add_paragraph('当前USD兑CNY价格为：%s'%(usd_price),style = 'ListBullet')
document.add_paragraph('当前USDT场外兑CNY价格为：%s'%(usdt_price),style = 'ListBullet')
document.add_paragraph('当前USDT场外溢价率为：%s'%(usdt_per),style = 'ListBullet')

#usdt流入交易所
document.add_paragraph('交易所USDT净流入量',style = 'ListBullet')
u_jinliuru_data = u_jinliuru_data.reset_index(drop=True)
usdt_netflow = u_jinliuru_data['value'][len(u_jinliuru_data)-1]
text = '昨日交易所USDT净流入量为：%s。'%(usdt_netflow)
document.add_paragraph(text,style = 'ListBullet')
# 添加图片，并指定宽度
document.add_picture('/root/xianhuo/usdt_netflow.png',width = Inches(5.25))


# ----- 重要链上指标
document.add_page_break()
#p = document.add_paragraph('This is a paragraph in new page.')
document.add_heading(u'3.重要链上指标',level = 1)

document.add_paragraph('短期行情',style = 'ListBullet')

res_df = res_df.reset_index(drop=True)
asopr_v = res_df['aSOPR'][len(res_df)-1]
sopr_7v = res_df['7MA aSOPR'][len(res_df)-1]


document.add_paragraph('aSOPR',style = 'ListBullet')
document.add_paragraph('当aSOPR值大于1时，说明BTC全网持有者总体处于盈利状态，当其小于1时，说明总体处于亏损状态，在目前市场状态下是看涨信号,可以少量现货进入。',style = 'ListBullet')
document.add_paragraph('昨日收盘值为：%s'%(asopr_v),style = 'ListBullet')
document.add_picture('/root/xianhuo/aSOPR.png',width = Inches(5.25))


document.add_paragraph('7MA aSOPR',style = 'ListBullet')
document.add_paragraph('当7MA aSOPR值大于1时，说明BTC全网持有者总体处于盈利状态，当其小于1时，说明总体处于亏损状态，在目前市场状态下是可以开启定投。',style = 'ListBullet')
document.add_paragraph('昨日收盘值为：%s'%(sopr_7v),style = 'ListBullet')
document.add_picture('/root/xianhuo/7MA_aSOPR.png',width = Inches(5.25))

document.add_paragraph('BTC RISK INDEX',style = 'ListBullet')
risk_data = risk_data.reset_index(drop=True)
risk = risk_data['value'][len(risk_data)-1]
document.add_paragraph('昨日收盘值为：%s'%(risk),style = 'ListBullet')
document.add_picture('/root/xianhuo/risk_index.png',width = Inches(5.25))

document.add_paragraph('Entity-Adjusted STH-NUPL',style = 'ListBullet')
nupl_data = nupl_data.reset_index(drop=True)
nupl = nupl_data['value'][len(nupl_data)-1]
document.add_paragraph('昨日收盘值为：%s'%(nupl),style = 'ListBullet')
document.add_picture('/root/xianhuo/sth_nupl.png',width = Inches(5.25))



document.add_paragraph('周期行情',style = 'ListBullet')

pm_v = res_df['Puell Multiple'][len(res_df)-1]
mvrv_v = res_df['MVRV Z-Score'][len(res_df)-1]
supply_v = res_df['Percent Supply in Profit'][len(res_df)-1]
rhodl_v = res_df['RHODL Ratio'][len(res_df)-1]


document.add_paragraph('MVRV Z-Score',style = 'ListBullet')
document.add_paragraph('当MVRV Z-Score值大于7时，是牛顶信号，当其小于0时，是熊底信号。',style = 'ListBullet')      
document.add_paragraph('昨日收盘值为：%s'%(mvrv_v),style = 'ListBullet')
document.add_picture('/root/xianhuo/MVRV_Z_Score.png',width = Inches(5.25))

document.add_paragraph('Puell Multiple',style = 'ListBullet')
document.add_paragraph('当Puell Multiple值大于4时，是牛顶信号，当其小于0.5时，是熊底信号。',style = 'ListBullet')
document.add_paragraph('昨日收盘值为：%s'%(pm_v),style = 'ListBullet')
document.add_picture('/root/xianhuo/Puell.png',width = Inches(5.25))

document.add_paragraph('Percent Supply in Profit',style = 'ListBullet')
document.add_paragraph('当Percent Supply in Profit值大于0.95时，是牛顶信号，当其小于0.5时，是熊底信号。',style = 'ListBullet')
document.add_paragraph('昨日收盘值为：%s'%(supply_v),style = 'ListBullet')
document.add_picture('/root/xianhuo/Percent_Supply.png',width = Inches(5.25))

document.add_paragraph('RHODL Ratio',style = 'ListBullet')
document.add_paragraph('当RHODL Ratio值大于49000时，是牛顶信号，当其小于350时，是熊底信号。',style = 'ListBullet')
document.add_paragraph('昨日收盘值为：%s'%(rhodl_v),style = 'ListBullet')
document.add_picture('/root/xianhuo/RHODL.png',width = Inches(5.25))

# -----  黑天鹅事件
document.add_page_break()
#p = document.add_paragraph('This is a paragraph in new page.')
document.add_heading(u'4.黑天鹅事件监控',level = 1)

document.add_paragraph(content_us,style = 'ListBullet')
document.add_paragraph(content_gox,style = 'ListBullet')
# 添加图片，并指定宽度
document.add_picture('/root/xianhuo/US_MT.png',width = Inches(5.25))

document.add_paragraph('今年美国政府转移BTC时间',style = 'ListBullet')
t = document.add_table(rows=1, cols=3) # 插入表格，先将表头写好，参数：rows:行，cols:列
hdr_cells = t.rows[0].cells
hdr_cells[0].text = '时间' # 表头
hdr_cells[1].text = '余额'# 表头
hdr_cells[2].text = '前一天余额'# 表头

for d in last_data_us.values.tolist(): # 
    print(d)
    for date,v1,v2 in [d]: # 读取每一行内容
        row_cells = t.add_row().cells # 读到一行就在word的表格中插入一行
        row_cells[0].text = str(date) 
        row_cells[1].text = str(v1)
        row_cells[2].text = str(v2)

# ----- 重要链上指标
sub_dev_data = dev_data[-7:]
sub_dev_data = sub_dev_data.reset_index(drop=True)
document.add_page_break()
#p = document.add_paragraph('This is a paragraph in new page.')
document.add_heading(u'5.衍生品指标',level = 1)

document.add_paragraph('Deribit比特币期权 25 Delta Skew（7日均值）',style = 'ListBullet')
document.add_paragraph('25 Delta Skew与比特币价格呈现明显的负相关关系，该值大说明比特币价格处于相对低位，该值小说明价格处于相对高位，要注意即使比特币价格在一直下跌，但是该值也在下跌，不是抄底的时刻。',style = 'ListBullet')
document.add_picture('/root/xianhuo/btc_options.png',width = Inches(5.25))

document.add_paragraph('比特币期权近7日值',style = 'ListBullet')
t = document.add_table(rows=1, cols=2) # 插入表格，先将表头写好，参数：rows:行，cols:列
hdr_cells = t.rows[0].cells
hdr_cells[0].text = '时间' # 表头
hdr_cells[1].text = '当日options值'# 表头

for d in sub_dev_data.values.tolist(): # 
    print(d)
    for date,v1,v2 in [d]: # 读取每一行内容
        row_cells = t.add_row().cells # 读到一行就在word的表格中插入一行
        row_cells[0].text = str(date) 
        row_cells[1].text = str(v1)

document.add_paragraph('永续合约总量（7日均值）',style = 'ListBullet')
document.add_paragraph('永续合约总量处于阶段性低值的时候，说明价格下跌已经引不起交易者兴趣，这时候是比特币短期现货定投的开始位置。',style = 'ListBullet')
document.add_picture('/root/xianhuo/btc_futures.png',width = Inches(5.25))

document.add_paragraph('永续合约总量近7日值',style = 'ListBullet')
t = document.add_table(rows=1, cols=2) # 插入表格，先将表头写好，参数：rows:行，cols:列
hdr_cells = t.rows[0].cells
hdr_cells[0].text = '时间' # 表头
hdr_cells[1].text = '永续合约当日总量'# 表头

for d in sub_dev_data.values.tolist(): # 
    print(d)
    for date,v1,v2 in [d]: # 读取每一行内容
        row_cells = t.add_row().cells # 读到一行就在word的表格中插入一行
        row_cells[0].text = str(date) 
        row_cells[1].text = str(v2)

# ===============================================================附录===================================================================
document.add_page_break()
#p = document.add_paragraph('This is a paragraph in new page.')
document.add_heading(u'附录',level = 1)

per_1 = combine_data['per'][len(combine_data)-1]
document.add_paragraph('ETH Price/BTC Price',style = 'ListBullet')
document.add_paragraph('当前该比值为：%s'%(per_1),style = 'ListBullet')
document.add_picture('/root/xianhuo/eth_btc.png',width = Inches(5.25))

doc_name = '%s日BTC链上数据一览'%(str(date_now)[0:10]) + '.doc'

# 保存文档
document.save(doc_name)
print('success')

import telegram

date_now = datetime.datetime.utcnow()
doc_name = '%s日BTC链上数据一览'%(str(date_now)[0:10]) + '.doc'


text_0 = '【综合分析】：经过近几日的价格上涨，有持续上升趋势，由于前期已经定投完成，所以继续观察走势，不操作。'
text_1 = '【综合分析】：可以尝试小资金进行抄底，虽然小资金，但是也要分批,因为价格下跌会有一个惯性趋势。'
text_2 = '【综合分析】：BTC价格持续下跌，前期小资金抄底博反弹失效，目前已经进入到了定投阶段，可以稍微加大资金量，但不能ALL IN，防止黑天鹅事件。'
text_3 = '【综合分析】：经过近几日的价格反弹，BTC已经暂时走出底部区域，前期抄底或定投的也可以分批套现。'

if asopr_v >=1 and sopr_7v >=1:
    text_4 = text_3
elif asopr_v < 1 and sopr_7v >=1:
    text_4 = text_1
elif asopr_v >= 1 and sopr_7v <1:
    text_4 = text_0
elif asopr_v < 1 and sopr_7v <1:
    text_4 = text_2
else:
    text_4 = 'unknow'

tg_doc_name = '/root/xianhuo/' + str(doc_name)
print(tg_doc_name)
bot = telegram.Bot(token='6361430672:AAG2qr7zuFQkcQb13Xtud2q8KksonuTNVN4')
bot.sendDocument(chat_id='-1001920263299', document=open(tg_doc_name, 'rb'),message_thread_id=5) #链上数据分享
bot.sendMessage(chat_id='-1001920263299', text = text_4,message_thread_id=5) #链上数据分享


