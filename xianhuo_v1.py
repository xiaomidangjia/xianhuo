#-*- conding:utf-8 -*-
import json
import requests
import pandas as pd
import time
import numpy as np
import os
import re
import datetime
from requests import Request,Session
from matplotlib import font_manager as fm, rcParams
from requests.exceptions import ConnectionError,Timeout,TooManyRedirects
import matplotlib.pyplot as plt
import seaborn as sns
import hashlib
import random
from PIL import Image, ImageDraw, ImageFont
import cv2
from watermarker.marker import add_mark

def jigsaw(imgs, direction, gap=0):
    imgs = [Image.fromarray(img) for img in imgs]
    w, h = imgs[0].size
    if direction == "horizontal":
        result = Image.new(imgs[0].mode, ((w+gap)*len(imgs)-gap, h))
        for i, img in enumerate(imgs):
            result.paste(img, box=((w+gap)*i, 0))
    elif direction == "vertical":
        result = Image.new(imgs[0].mode, (w, (h+gap)*len(imgs)-gap))
        for i, img in enumerate(imgs):
            result.paste(img, box=(0, (h+gap)*i))
    else:
        raise ValueError("The direction parameter has only two options: horizontal and vertical")
    return np.array(result)
# ======= 正式开始执行
prop = fm.FontProperties(fname='/root/xianhuo/SimHei.ttf')
#pip install jojo-office
#import office
#office.open_file("output.docx").save("output.pdf")

# ----- 基础btc数据
url_address = [ 'https://api.glassnode.com/v1/metrics/market/price_usd_ohlc']
url_name = ['k_fold']
# insert your API key here
API_KEY = '26BLocpWTcSU7sgqDdKzMHMpJDm'
data_list = []
for num in range(len(url_name)):
    print(num)
    addr = url_address[num]
    name = url_name[num]
    # make API request
    res_addr = requests.get(addr,params={'a': 'BTC', 'api_key': API_KEY})
    # convert to pandas dataframe
    ins = pd.read_json(res_addr.text, convert_dates=['t'])
    #ins.to_csv('test.csv')
    #print(ins['o'])
    ins['date'] =  ins['t']
    ins['value'] =  ins['o']
    ins = ins[['date','value']]
    data_list.append(ins)
result_data = data_list[0][['date']]
for i in range(len(data_list)):
    df = data_list[i]
    result_data = result_data.merge(df,how='left',on='date')
#last_data = result_data[(result_data.date>='2016-01-01') & (result_data.date<='2020-01-01')]
last_data = result_data[(result_data.date>='2013-01-01')]
last_data = last_data.sort_values(by=['date'])
last_data = last_data.reset_index(drop=True)
print(type(last_data))
date = []
open_p = []
close_p = []
high_p = []
low_p = []
for i in range(len(last_data)):
    date.append(last_data['date'][i])
    open_p.append(last_data['value'][i]['o'])
    close_p.append(last_data['value'][i]['c'])
    high_p.append(last_data['value'][i]['h'])
    low_p.append(last_data['value'][i]['l'])
#res_data = pd.DataFrame({'date':date,'open':open_p,'close':close_p,'high':high_p,'low':low_p})
price_data = pd.DataFrame({'date':date,'close':close_p})
price_data = price_data.sort_values(by=['date'])
price_data = price_data.reset_index(drop=True)


# =======================================================================btc链上数据===================================================================

# ----- btc市值占比
url_address = [ 'https://api.glassnode.com/v1/metrics/market/btc_dominance']
url_name = ['k_fold']
# insert your API key here
API_KEY = '26BLocpWTcSU7sgqDdKzMHMpJDm'
data_list = []
for num in range(len(url_name)):
    print(num)
    addr = url_address[num]
    name = url_name[num]
    # make API request
    res_addr = requests.get(addr,params={'a': 'BTC', 'api_key': API_KEY})
    # convert to pandas dataframe
    ins = pd.read_json(res_addr.text, convert_dates=['t'])
    #ins.to_csv('test.csv')
    #print(ins['o'])
    ins['date'] =  ins['t']
    ins['value'] =  ins['v']
    ins = ins[['date','value']]
    data_list.append(ins)
result_data = data_list[0][['date']]
for i in range(len(data_list)):
    df = data_list[i]
    result_data = result_data.merge(df,how='left',on='date')
#last_data = result_data[(result_data.date>='2016-01-01') & (result_data.date<='2020-01-01')]
domain_data = result_data[(result_data.date>='2013-01-01')]
domain_data = domain_data.sort_values(by=['date'])
domain_data = domain_data.reset_index(drop=True)
domain_data = domain_data.merge(price_data,how='left',on=['date'])


f, axes = plt.subplots(figsize=(20, 10))
axes_fu = axes.twinx()
sns.lineplot(x="date", y="value",color='red',data=domain_data, ax=axes_fu)
sns.lineplot(x="date", y="close",color='black', data=domain_data, ax=axes)
axes.tick_params(labelsize=10)
plt.title('BTC Dominance', fontsize=10) 
axes.legend(loc='upper left', fontsize=5)
axes.set_xlabel('时间',fontsize=14,fontproperties=prop)
axes.set_ylabel("BTC price",fontsize=10)
axes_fu.set_ylabel("BTC Dominance",fontsize=10)

#plt.show()
plt.savefig('btc_dominance.png')
plt.close()
add_mark(file = "btc_dominance.png", out = "out",mark = "MMC研究猿卡森出品", opacity=0.2, angle=30, space=30)


# ----- 矿工
url_address = ['https://api.glassnode.com/v1/metrics/transactions/transfers_volume_miners_to_exchanges']
url_name = ['k_fold']
# insert your API key here
API_KEY = '26BLocpWTcSU7sgqDdKzMHMpJDm'
data_list = []
for num in range(len(url_name)):
    print(num)
    addr = url_address[num]
    name = url_name[num]
    # make API request
    res_addr = requests.get(addr,params={'a': 'BTC', 'api_key': API_KEY})
    # convert to pandas dataframe
    ins = pd.read_json(res_addr.text, convert_dates=['t'])
    #ins.to_csv('test.csv')
    #print(ins['o'])
    ins['date'] =  ins['t']
    ins['value'] =  ins['v']
    ins = ins[['date','value']]
    data_list.append(ins)
result_data = data_list[0][['date']]
for i in range(len(data_list)):
    df = data_list[i]
    result_data = result_data.merge(df,how='left',on='date')
#last_data = result_data[(result_data.date>='2016-01-01') & (result_data.date<='2020-01-01')]
miner_data = result_data[(result_data.date>='2013-01-01')]
miner_data = miner_data.sort_values(by=['date'])
miner_data = miner_data.reset_index(drop=True)


date = []
miner_raw = []
for j in range(20,len(miner_data)+1):
    ins = miner_data[j-14:j]
    ins = ins.sort_values(by='date')
    ins = ins.reset_index(drop=True)
    date.append(ins['date'][13])
    miner_raw.append(ins['value'][13])
miner_data_1 = pd.DataFrame({'date':date,'miner_raw':miner_raw})
miner_data_1 = miner_data_1[(miner_data_1.date>='2019-01-01')]


miner_data = miner_data_1.merge(price_data,how='left',on=['date'])


f, axes = plt.subplots(figsize=(20, 10))
axes_fu = axes.twinx()
sns.lineplot(x="date", y="miner_raw",color='red',data=miner_data, ax=axes_fu)
sns.lineplot(x="date", y="close",color='black', data=miner_data, ax=axes)
axes.tick_params(labelsize=10)
plt.title('Miner To Exhanges', fontsize=10) 
axes.legend(loc='upper left', fontsize=5)
axes.set_xlabel('时间',fontsize=14,fontproperties=prop)
axes.set_ylabel("BTC price",fontsize=10)
axes_fu.set_ylabel("BTC Volume",fontsize=10)

#plt.show()
plt.savefig('btc_miner_v.png')
plt.close()

add_mark(file = "btc_miner_v.png", out = "out",mark = "MMC研究猿卡森出品", opacity=0.2, angle=30, space=30)



# ----- 巨鲸
url_address = [ 'https://api.glassnode.com/v1/metrics/transactions/transfers_volume_whales_to_exchanges_sum']
url_name = ['k_fold']
# insert your API key here
API_KEY = '26BLocpWTcSU7sgqDdKzMHMpJDm'
data_list = []
for num in range(len(url_name)):
    print(num)
    addr = url_address[num]
    name = url_name[num]
    # make API request
    res_addr = requests.get(addr,params={'a': 'BTC', 'api_key': API_KEY})
    # convert to pandas dataframe
    ins = pd.read_json(res_addr.text, convert_dates=['t'])
    #ins.to_csv('test.csv')
    #print(ins['o'])
    ins['date'] =  ins['t']
    ins['value'] =  ins['v']
    ins = ins[['date','value']]
    data_list.append(ins)
result_data = data_list[0][['date']]
for i in range(len(data_list)):
    df = data_list[i]
    result_data = result_data.merge(df,how='left',on='date')
#last_data = result_data[(result_data.date>='2016-01-01') & (result_data.date<='2020-01-01')]
whale_data = result_data[(result_data.date>='2013-01-01')]
whale_data = whale_data.sort_values(by=['date'])
whale_data = whale_data.reset_index(drop=True)
whale_data = whale_data.merge(price_data,how='left',on=['date'])


date = []
whale_raw = []
for j in range(30,len(whale_data)+1):
    ins = whale_data[j-30:j]
    ins = ins.sort_values(by='date')
    ins = ins.reset_index(drop=True)
    date.append(ins['date'][29])
    whale_raw.append(ins['value'][29])
whale_data_1 = pd.DataFrame({'date':date,'whale_raw':whale_raw})
whale_data_1 = whale_data_1[(whale_data_1.date>='2019-01-01')]


whale_data = whale_data_1.merge(price_data,how='left',on=['date'])


f, axes = plt.subplots(figsize=(20, 10))
axes_fu = axes.twinx()
sns.lineplot(x="date", y="whale_raw",color='red',data=whale_data, ax=axes_fu)
sns.lineplot(x="date", y="close",color='black', data=whale_data, ax=axes)
axes.tick_params(labelsize=10)
plt.title('Whale To Exhanges', fontsize=10) 
axes.legend(loc='upper left', fontsize=5)
axes.set_xlabel('时间',fontsize=14,fontproperties=prop)
axes.set_ylabel("BTC price",fontsize=10)
axes_fu.set_ylabel("Whale Volume",fontsize=10)

#plt.show()
plt.savefig('btc_whale_v.png')
plt.close()
add_mark(file = "btc_whale_v.png", out = "out",mark = "MMC研究猿卡森出品", opacity=0.2, angle=30, space=30)




# ---- btc净流入交易所

url_address = [ 'https://api.glassnode.com/v1/metrics/transactions/transfers_volume_exchanges_net']
url_name = ['k_fold']
# insert your API key here
API_KEY = '26BLocpWTcSU7sgqDdKzMHMpJDm'
data_list = []
for num in range(len(url_name)):
    print(num)
    addr = url_address[num]
    name = url_name[num]
    # make API request
    res_addr = requests.get(addr,params={'a': 'BTC', 'api_key': API_KEY})
    # convert to pandas dataframe
    ins = pd.read_json(res_addr.text, convert_dates=['t'])
    #ins.to_csv('test.csv')
    #print(ins['o'])
    ins['date'] =  ins['t']
    ins['value'] =  ins['v']
    ins = ins[['date','value']]
    data_list.append(ins)
result_data = data_list[0][['date']]
for i in range(len(data_list)):
    df = data_list[i]
    result_data = result_data.merge(df,how='left',on='date')
#last_data = result_data[(result_data.date>='2016-01-01') & (result_data.date<='2020-01-01')]
jinliuru_data = result_data[(result_data.date>='2013-01-01')]
jinliuru_data = jinliuru_data.sort_values(by=['date'])
jinliuru_data = jinliuru_data.reset_index(drop=True)
jinliuru_data = jinliuru_data.merge(price_data,how='left',on=['date'])
jinliuru_data = jinliuru_data[-120:]
f, axes = plt.subplots(figsize=(20, 10))
sns.barplot(x="date", y="value",data=jinliuru_data, ax=axes)
#sns.lineplot(x="date", y="close",color='black', data=jinliuru_data, ax=axes)
axes.tick_params(labelsize=10)
plt.title('BTC Netflow Exchanges', fontsize=10) 
axes.legend(loc='upper left', fontsize=5)
axes.set_xlabel('时间',fontsize=14,fontproperties=prop)
axes.set_ylabel("Netflow",fontsize=10)
#plt.show()
plt.savefig('btc_netflow.png')
plt.close()
add_mark(file = "btc_netflow.png", out = "out",mark = "MMC研究猿卡森出品", opacity=0.2, angle=30, space=30)

f, axes = plt.subplots(figsize=(20, 10))
#sns.barplot(x="date", y="value",data=jinliuru_data, ax=axes)
sns.lineplot(x="date", y="close",color='black', data=jinliuru_data, ax=axes)
axes.tick_params(labelsize=10)
plt.title('BTC Price', fontsize=10) 
axes.legend(loc='upper left', fontsize=5)
axes.set_xlabel('时间',fontsize=14,fontproperties=prop)
axes.set_ylabel("Price",fontsize=10)
#plt.show()
plt.savefig('btc_price.png')
plt.close()

img1 = cv2.imread("btc_price.png")
img2 = cv2.imread("btc_netflow.png")

img = jigsaw([img1, img2],direction="vertical")
name = 'btc_netflow_price.png'
cv2.imwrite(name, img)
add_mark(file = "btc_netflow_price.png", out = "out",mark = "MMC研究猿卡森出品", opacity=0.2, angle=30, space=30)



# ----- usdt
url_address = [ 'https://api.glassnode.com/v1/metrics/transactions/transfers_volume_exchanges_net']
url_name = ['k_fold']
# insert your API key here
API_KEY = '26BLocpWTcSU7sgqDdKzMHMpJDm'
data_list = []
for num in range(len(url_name)):
    print(num)
    addr = url_address[num]
    name = url_name[num]
    # make API request
    res_addr = requests.get(addr,params={'a': 'USDT', 'api_key': API_KEY})
    # convert to pandas dataframe
    ins = pd.read_json(res_addr.text, convert_dates=['t'])
    #ins.to_csv('test.csv')
    #print(ins['o'])
    ins['date'] =  ins['t']
    ins['value'] =  ins['v']
    ins = ins[['date','value']]
    data_list.append(ins)
result_data = data_list[0][['date']]
for i in range(len(data_list)):
    df = data_list[i]
    result_data = result_data.merge(df,how='left',on='date')
#last_data = result_data[(result_data.date>='2016-01-01') & (result_data.date<='2020-01-01')]
u_jinliuru_data = result_data[(result_data.date>='2013-01-01')]
u_jinliuru_data = u_jinliuru_data.sort_values(by=['date'])
u_jinliuru_data = u_jinliuru_data.reset_index(drop=True)
u_jinliuru_data = u_jinliuru_data.merge(price_data,how='left',on=['date'])
u_jinliuru_data = u_jinliuru_data[-120:]
f, axes = plt.subplots(figsize=(20, 10))
sns.barplot(x="date", y="value",data=u_jinliuru_data, ax=axes)
#sns.lineplot(x="date", y="close",color='black', data=jinliuru_data, ax=axes)
axes.tick_params(labelsize=10)
plt.title('USDT Netflow Exchanges', fontsize=10) 
axes.legend(loc='upper left', fontsize=5)
axes.set_xlabel('时间',fontsize=14,fontproperties=prop)
axes.set_ylabel("Netflow",fontsize=10)
#plt.show()
plt.savefig('usdt_netflow.png')
plt.close()
f, axes = plt.subplots(figsize=(20, 10))
#sns.barplot(x="date", y="value",data=jinliuru_data, ax=axes)
sns.lineplot(x="date", y="close",color='black', data=u_jinliuru_data, ax=axes)
axes.tick_params(labelsize=10)
plt.title('BTC Price', fontsize=10) 
axes.legend(loc='upper left', fontsize=5)
axes.set_xlabel('时间',fontsize=14,fontproperties=prop)
axes.set_ylabel("Price",fontsize=10)
#plt.show()
plt.savefig('btc_price_1.png')
plt.close()


img1 = cv2.imread("btc_price_1.png")
img2 = cv2.imread("usdt_netflow.png")
img = jigsaw([img1, img2],direction="vertical")
name = 'usdt_netflow_price.png'
cv2.imwrite(name, img)
add_mark(file = "usdt_netflow_price.png", out = "out",mark = "MMC研究猿卡森出品", opacity=0.2, angle=30, space=30)

# btc_risk_index
url_address = [ 'https://api.glassnode.com/v1/metrics/signals/btc_risk_index']
url_name = ['k_fold']
# insert your API key here
API_KEY = '26BLocpWTcSU7sgqDdKzMHMpJDm'
data_list = []
for num in range(len(url_name)):
    print(num)
    addr = url_address[num]
    name = url_name[num]
    # make API request
    res_addr = requests.get(addr,params={'a': 'BTC', 'api_key': API_KEY})
    # convert to pandas dataframe
    ins = pd.read_json(res_addr.text, convert_dates=['t'])
    #ins.to_csv('test.csv')
    #print(ins['o'])
    ins['date'] =  ins['t']
    ins['value'] =  ins['v']
    ins = ins[['date','value']]
    data_list.append(ins)
result_data = data_list[0][['date']]
for i in range(len(data_list)):
    df = data_list[i]
    result_data = result_data.merge(df,how='left',on='date')
#last_data = result_data[(result_data.date>='2016-01-01') & (result_data.date<='2020-01-01')]
risk_data = result_data[(result_data.date>='2013-01-01')]
risk_data = risk_data.sort_values(by=['date'])
risk_data = risk_data.reset_index(drop=True)
risk_data = risk_data.merge(price_data,how='left',on=['date'])
risk_data = risk_data[risk_data.date>='2023-01-01']
risk_data['x1'] = 0.25
f, axes = plt.subplots(figsize=(20, 10))
axes_fu = axes.twinx()
sns.lineplot(x="date", y="value",data=risk_data, ax=axes)
sns.lineplot(x="date", y="x1",data=risk_data, ax=axes)
sns.lineplot(x="date", y="close",color='black', data=risk_data, ax=axes_fu)
axes.tick_params(labelsize=10)
plt.title('BTC RISK INDEX', fontsize=10) 
axes.legend(loc='upper left', fontsize=5)
axes.set_xlabel('时间',fontsize=14,fontproperties=prop)
axes.set_ylabel("INDEX",fontsize=10)
#plt.show()
plt.savefig('risk_index.png')
plt.close()
add_mark(file = "risk_index.png", out = "out",mark = "MMC研究猿卡森出品", opacity=0.2, angle=30, space=30)


# ----- 基础btc数据
url_address = [ 'https://api.glassnode.com/v1/metrics/indicators/nupl_less_155_account_based']
url_name = ['k_fold']
# insert your API key here
API_KEY = '26BLocpWTcSU7sgqDdKzMHMpJDm'
data_list = []
for num in range(len(url_name)):
    print(num)
    addr = url_address[num]
    name = url_name[num]
    # make API request
    res_addr = requests.get(addr,params={'a': 'BTC', 'api_key': API_KEY})
    # convert to pandas dataframe
    ins = pd.read_json(res_addr.text, convert_dates=['t'])
    #ins.to_csv('test.csv')
    #print(ins['o'])
    ins['date'] =  ins['t']
    ins['value'] =  ins['v']
    ins = ins[['date','value']]
    data_list.append(ins)
result_data = data_list[0][['date']]
for i in range(len(data_list)):
    df = data_list[i]
    result_data = result_data.merge(df,how='left',on='date')
#last_data = result_data[(result_data.date>='2016-01-01') & (result_data.date<='2020-01-01')]
nupl_data = result_data[(result_data.date>='2013-01-01')]
nupl_data = nupl_data.sort_values(by=['date'])
nupl_data = nupl_data.reset_index(drop=True)
nupl_data = nupl_data.merge(price_data,how='left',on=['date'])
nupl_data = nupl_data[nupl_data.date>='2023-01-01']
nupl_data['x1'] = 0
nupl_data['x2'] = 0.25
f, axes = plt.subplots(figsize=(20, 10))
axes_fu = axes.twinx()
sns.lineplot(x="date", y="value",color='red',data=nupl_data, ax=axes)
sns.lineplot(x="date", y="x1",color='green',linestyle='--',data=nupl_data, ax=axes)
sns.lineplot(x="date", y="x2",color='green',linestyle='--',data=nupl_data, ax=axes)
sns.lineplot(x="date", y="close",color='black', data=nupl_data, ax=axes_fu)
axes.tick_params(labelsize=10)
plt.title('Entity-Adjusted STH-NUPL', fontsize=10) 
axes.legend(loc='upper left', fontsize=5)
axes.set_xlabel('时间',fontsize=14,fontproperties=prop)
axes.set_ylabel("STH-NUPL",fontsize=10)
#plt.show()
plt.savefig('sth_nupl.png')
plt.close()
add_mark(file = "sth_nupl.png", out = "out",mark = "MMC研究猿卡森出品", opacity=0.2, angle=30, space=30)

# =======================================================================黑天鹅地址监控===================================================================

url_address = [ 'https://api.glassnode.com/v1/metrics/distribution/balance_us_government','https://api.glassnode.com/v1/metrics/distribution/balance_mtgox_trustee','https://api.glassnode.com/v1/metrics/market/price_usd_close']
url_name = ['balance_us','balance_mtgox','Price']
# insert your API key here
API_KEY = '26BLocpWTcSU7sgqDdKzMHMpJDm'
data_list = []
for num in range(len(url_name)):
    print(num)
    addr = url_address[num]
    name = url_name[num]
    # make API request
    res_addr = requests.get(addr,params={'a': 'BTC', 'api_key': API_KEY})
    # convert to pandas dataframe
    ins = pd.read_json(res_addr.text, convert_dates=['t'])
    #ins.to_csv('test.csv')
    #print(ins['o'])
    ins['date'] =  ins['t']
    ins['value'] =  ins['v']
    ins = ins[['date','value']]
    data_list.append(ins)
result_data = data_list[0][['date']]
for i in range(len(data_list)):
    df = data_list[i]
    result_data = result_data.merge(df,how='left',on='date')
#last_data = result_data[(result_data.date>='2016-01-01') & (result_data.date<='2020-01-01')]
last_data = result_data[(result_data.date>='2018-01-01')]
last_data['log(BTC price)'] = np.log(last_data['value'])
last_data.rename(columns={'value_x':'Balance US','value_y':'Balance Mt.GOX'},inplace=True)
last_data = last_data.sort_values(by=['date'])
last_data = last_data.reset_index(drop=True)

last_data_us = last_data[['date','Balance US']]
last_data_us['nb'] = last_data_us['Balance US'].shift(-1)
last_data_us['nd'] = last_data_us['date'].shift(-1)
last_data_us = last_data_us[0:-1]
last_data_us['cha'] = last_data_us['nb'] - last_data_us['Balance US']
last_data_us = last_data_us[last_data_us.cha < 0]
last_data_us = last_data_us[['nd','nb','Balance US']]

us_pre_b = last_data['Balance US'][len(last_data)-2]
us_pre_t = last_data['Balance US'][len(last_data)-1]

gox_pre_b = last_data['Balance Mt.GOX'][len(last_data)-2]
gox_pre_t = last_data['Balance Mt.GOX'][len(last_data)-1]


change_us = us_pre_t - us_pre_b
if change_us == 0:
    content_us = '美国政府持有地址目前有%s个BTC，相比昨日无变化.'%(str(us_pre_t))
elif change_us < 0:
    content_us = '美国政府持有地址目前有%s个BTC，相比昨日减少%s个BTC.'%(str(us_pre_t),str(-change_us))
else:
    content_us = '美国政府持有地址目前有%s个BTC，相比昨日增加%s个BTC.'%(str(us_pre_t),str(change_us))
print(content_us)
change_gox = gox_pre_t - gox_pre_b
if change_gox == 0:
    content_gox = 'MT.GOX持有地址目前有%s个BTC，相比昨日无变化.'%(str(gox_pre_t))
elif change_gox < 0:
    content_gox = 'MT.GOX持有地址目前有%s个BTC，相比昨日减少%s个BTC.'%(str(gox_pre_t),str(-change_gox))
else:
    content_gox = 'MT.GOX持有地址目前有%s个BTC，相比昨日增加%s个BTC.'%(str(gox_pre_t),str(change_gox))
f, axes = plt.subplots(figsize=(20, 10))
axes_fu = axes.twinx()

sns.lineplot(x="date", y="log(BTC price)", data=last_data, color = 'black', linewidth=0.5,ci=95, ax=axes)
sns.lineplot(x="date", y="Balance US",color='red', linewidth=0.5,data=last_data, ci=95, ax=axes_fu)
sns.lineplot(x="date", y="Balance Mt.GOX",color='green', linewidth=0.5,data=last_data, ci=95, ax=axes_fu)
axes.set_title('US GOV和MT GOX持币地址BTC数量', fontsize=30,fontproperties=prop) 
axes.legend(loc='upper left', fontsize=10)
axes.set_xlabel('时间',fontsize=14,fontproperties=prop)
axes.set_ylabel("BTC价格",fontsize=14,fontproperties=prop)
axes_fu.set_ylabel("持有BTC总量",fontsize=14,fontproperties=prop)
plt.savefig('US_MT.png',  bbox_inches='tight')
plt.close()

add_mark(file = "US_MT.png", out = "out",mark = "MMC研究猿卡森出品", opacity=0.2, angle=30, space=30)



# =======================================================================大额转账===================================================================

zhuanzhang_df = pd.read_csv('/root/whale_alert_tocsv/res_alert.csv',encoding='gbk')
zhuanzhang_df['date'] = pd.to_datetime(zhuanzhang_df['date'])
#zhuanzhang_df['date'] = zhuanzhang_df['date'] - datetime.timedelta(hours=8)
date_now = datetime.datetime.utcnow()
#print(date_now)
date_before = pd.to_datetime(date_now) - datetime.timedelta(hours=6)
zhuanzhang_df = zhuanzhang_df[zhuanzhang_df.date>=date_before]
zhuanzhang_df = zhuanzhang_df.drop_duplicates()


# =======================================================================链上的4种稳定币===================================================================

url_address = [ 'https://api.glassnode.com/v1/metrics/supply/current']
url_name = ['supply_usdt']
# insert your API key here
API_KEY = '26BLocpWTcSU7sgqDdKzMHMpJDm'
data_list = []
for num in range(len(url_name)):
    print(num)
    addr = url_address[num]
    name = url_name[num]
    # make API request
    res_addr = requests.get(addr,params={'a': 'USDT', 'api_key': API_KEY})
    # convert to pandas dataframe
    ins = pd.read_json(res_addr.text, convert_dates=['t'])
    #ins.to_csv('test.csv')
    #print(ins['o'])
    ins['date'] =  ins['t']
    ins['value'] =  ins['v']
    ins = ins[['date','value']]
    data_list.append(ins)
result_data = data_list[0][['date']]
for i in range(len(data_list)):
    df = data_list[i]
    result_data = result_data.merge(df,how='left',on='date')
#last_data = result_data[(result_data.date>='2016-01-01') & (result_data.date<='2020-01-01')]
usdt_data = result_data[(result_data.date>='2018-01-01')]
usdt_data.rename(columns={'value':'supply_usdt'},inplace=True)
url_address = [ 'https://api.glassnode.com/v1/metrics/supply/current']
url_name = ['supply_usdc']
# insert your API key here
API_KEY = '26BLocpWTcSU7sgqDdKzMHMpJDm'
data_list = []
for num in range(len(url_name)):
    print(num)
    addr = url_address[num]
    name = url_name[num]
    # make API request
    res_addr = requests.get(addr,params={'a': 'USDC', 'api_key': API_KEY})
    # convert to pandas dataframe
    ins = pd.read_json(res_addr.text, convert_dates=['t'])
    #ins.to_csv('test.csv')
    #print(ins['o'])
    ins['date'] =  ins['t']
    ins['value'] =  ins['v']
    ins = ins[['date','value']]
    data_list.append(ins)
result_data = data_list[0][['date']]
for i in range(len(data_list)):
    df = data_list[i]
    result_data = result_data.merge(df,how='left',on='date')
#last_data = result_data[(result_data.date>='2016-01-01') & (result_data.date<='2020-01-01')]
usdc_data = result_data[(result_data.date>='2018-01-01')]
usdc_data.rename(columns={'value':'supply_usdc'},inplace=True)
url_address = [ 'https://api.glassnode.com/v1/metrics/supply/current']
url_name = ['supply_busd']
# insert your API key here
API_KEY = '26BLocpWTcSU7sgqDdKzMHMpJDm'
data_list = []
for num in range(len(url_name)):
    print(num)
    addr = url_address[num]
    name = url_name[num]
    # make API request
    res_addr = requests.get(addr,params={'a': 'BUSD', 'api_key': API_KEY})
    # convert to pandas dataframe
    ins = pd.read_json(res_addr.text, convert_dates=['t'])
    #ins.to_csv('test.csv')
    #print(ins['o'])
    ins['date'] =  ins['t']
    ins['value'] =  ins['v']
    ins = ins[['date','value']]
    data_list.append(ins)
result_data = data_list[0][['date']]
for i in range(len(data_list)):
    df = data_list[i]
    result_data = result_data.merge(df,how='left',on='date')
#last_data = result_data[(result_data.date>='2016-01-01') & (result_data.date<='2020-01-01')]
busd_data = result_data[(result_data.date>='2018-01-01')]
busd_data.rename(columns={'value':'supply_busd'},inplace=True)
url_address = [ 'https://api.glassnode.com/v1/metrics/supply/current']
url_name = ['supply_tusd']
# insert your API key here
API_KEY = '26BLocpWTcSU7sgqDdKzMHMpJDm'
data_list = []
for num in range(len(url_name)):
    print(num)
    addr = url_address[num]
    name = url_name[num]
    # make API request
    res_addr = requests.get(addr,params={'a': 'TUSD', 'api_key': API_KEY})
    # convert to pandas dataframe
    ins = pd.read_json(res_addr.text, convert_dates=['t'])
    #ins.to_csv('test.csv')
    #print(ins['o'])
    ins['date'] =  ins['t']
    ins['value'] =  ins['v']
    ins = ins[['date','value']]
    data_list.append(ins)
result_data = data_list[0][['date']]
for i in range(len(data_list)):
    df = data_list[i]
    result_data = result_data.merge(df,how='left',on='date')
#last_data = result_data[(result_data.date>='2016-01-01') & (result_data.date<='2020-01-01')]
tusd_data = result_data[(result_data.date>='2018-01-01')]
tusd_data.rename(columns={'value':'supply_tusd'},inplace=True)
all_stable_coin = usdt_data.merge(usdc_data,how='left',on=['date']).merge(busd_data,how='left',on=['date']).merge(tusd_data,how='left',on=['date'])
all_stable_coin = all_stable_coin.fillna(0)
all_stable_coin['all'] = all_stable_coin['supply_usdt'] + all_stable_coin['supply_usdc'] + all_stable_coin['supply_busd'] + all_stable_coin['supply_tusd']
all_stable_coin = all_stable_coin.merge(price_data,how='left',on=['date'])
all_stable_coin = all_stable_coin.sort_values(by='date')
all_stable_coin = all_stable_coin.reset_index(drop=True)
all_stable_coin['Toal Supply'] = all_stable_coin['all']/100000000
all_stable_coin = all_stable_coin[all_stable_coin.date>='2022-01-01']
f, axes = plt.subplots(figsize=(20, 10))
axes_fu = axes.twinx()

sns.lineplot(x="date", y="close",color='black', linewidth=0.5,data=all_stable_coin, ci=95, ax=axes)
sns.lineplot(x="date", y="Toal Supply",color='blue', linewidth=0.5,data=all_stable_coin, ci=95, ax=axes_fu)
plt.title('USDT/USDC/BUSD/TUSD链上总发行量', fontsize=20,fontproperties=prop) 
axes.set_xlabel('时间',fontsize=14,fontproperties=prop)
axes.set_ylabel("BTC价格",fontsize=14,fontproperties=prop)
axes_fu.set_ylabel("发行总量",fontsize=14,fontproperties=prop)
plt.savefig('BTC价格-稳定币发行总量.png',  bbox_inches='tight')
plt.close()

add_mark(file = "BTC价格-稳定币发行总量.png", out = "out",mark = "MMC研究猿卡森出品", opacity=0.2, angle=30, space=30)


# =======================================================================美国经济数据===================================================================

us_eco_df = pd.read_csv('/root/xianhuo/us_ec_data.csv')
us_eco_df['date'] = us_eco_df['下一次公布日期']
sub_us_eco_df = us_eco_df[us_eco_df.date==str(date_now)[0:10]]


url_address = ['https://api.glassnode.com/v1/metrics/market/price_usd_close']
url_name = ['Price']
# insert your API key here
API_KEY = '26BLocpWTcSU7sgqDdKzMHMpJDm'
data_list = []
for num in range(len(url_name)):
    print(num)
    addr = url_address[num]
    name = url_name[num]
    # make API request
    res_addr = requests.get(addr,params={'a': 'BTC', 'api_key': API_KEY})
    # convert to pandas dataframe
    ins = pd.read_json(res_addr.text, convert_dates=['t'])
    ins['date'] =  ins['t']
    ins[name] =  ins['v']
    ins = ins[['date',name]]
    data_list.append(ins)
result_data = data_list[0][['date']]
for i in range(len(data_list)):
    df = data_list[i]
    result_data = result_data.merge(df,how='left',on='date')
#last_data = result_data[(result_data.date>='2016-01-01') & (result_data.date<='2020-01-01')]
last_data = result_data[(result_data.date>='2010-01-01')]
from dateutil.relativedelta import relativedelta 
#last_data['new_date'] = last_data['date'].apply(lambda x:x + relativedelta(years=1))
last_data = last_data.sort_values(by=['date'])
last_data = last_data.reset_index(drop=True)
date = []
price_raw = []
price_ma120 = []
price_ma200 = []
price_ma4y = []
for j in range(len(last_data)-1824):
    ins = last_data[j:j+1825]
    ins = ins.sort_values(by='date')
    ins = ins.reset_index(drop=True)
    date.append(ins['date'][1824])
    price_raw.append(ins['Price'][1824])
    price_ma4y.append(np.mean(ins['Price'][-1459:]))
    price_ma200.append(np.mean(ins['Price'][-199:]))
    price_ma120.append(np.mean(ins['Price'][-119:]))
jun_df = pd.DataFrame({'date':date,'price_raw':price_raw,'price_ma120':price_ma120,'price_ma200':price_ma200,'price_ma4y':price_ma4y})
jun_df = jun_df[(jun_df.date>='2018-12-01')]
f, axes = plt.subplots(figsize=(20, 10))
axes_fu = axes.twinx()

sns.lineplot(x="date", y="price_raw", data=jun_df, color = 'black', linewidth=0.5,ci=95, ax=axes)
sns.lineplot(x="date", y="price_ma120",color='red', linewidth=0.5,data=jun_df, ci=95, ax=axes)
sns.lineplot(x="date", y="price_ma200",color='green', linewidth=0.5,data=jun_df, ci=95, ax=axes)
sns.lineplot(x="date", y="price_ma4y",color='blue', linewidth=0.5,data=jun_df, ci=95, ax=axes)
plt.title('BTC价格-120D/200D/4Y线', fontsize=20,fontproperties=prop) 
axes.set_xlabel('时间',fontsize=14,fontproperties=prop)
axes.set_ylabel("BTC价格",fontsize=14,fontproperties=prop)
plt.savefig('BTC价格120D.png',  bbox_inches='tight')
plt.close()

add_mark(file = "BTC价格120D.png", out = "out",mark = "MMC研究猿卡森出品", opacity=0.2, angle=30, space=30)





# =======================================================================未来24小时价格预测==================================================================

date_now = datetime.datetime.utcnow()
date_before = pd.to_datetime(date_now) - datetime.timedelta(days=8)
# 读取原始数据
raw_data = pd.read_csv('/root/usdt/usdt.csv')
raw_data = raw_data[2:]
raw_data['date'] = raw_data['date'].apply(lambda x: str(x)[6:10] + '/' + str(x)[3:5] + '/' + str(x)[0:2] + ' ' + str(x)[11:19])
raw_data['date'] = pd.to_datetime(raw_data['date'])
raw_data['date'] = raw_data['date'] - datetime.timedelta(hours=8)
#取最近7天的数据
raw_data = raw_data[(raw_data.date >= date_before) & (raw_data.date < date_now)]
raw_data['date_dd'] = raw_data['date'].apply(lambda x: str(x)[0:10])
raw_data['date_hh'] = raw_data['date'].apply(lambda x: str(x)[11:13])
raw_data['per'] = raw_data['usdt']/raw_data['usd']
usdt_data = raw_data
raw_data = raw_data[['date_dd','date_hh','per']]
raw_data = raw_data.drop_duplicates()
raw_data = raw_data.groupby(['date_dd','date_hh'],as_index=False)['per'].mean()

raw_data = raw_data.sort_values(by=['date_dd','date_hh'])
raw_data = raw_data.reset_index(drop=True)
raw_data['per_1'] = raw_data['per'].shift(1)
raw_data['per_2'] = raw_data['per'].shift(2)
raw_data = raw_data.fillna(0)
per_last = []
for i in range(len(raw_data)):
    #print(i,raw_data['per'][i])
    if raw_data['per'][i] > 0:
        per_last.append(raw_data['per'][i])
    elif raw_data['per'][i] == 0:
        per_last.append(raw_data['per_1'][i])
    elif raw_data['per_1'][i] == 0:
        per_last.append(raw_data['per_2'][i])
    else:
        per_last.append(raw_data['per'][i])
raw_data['per_last'] = per_last
raw_data = raw_data[['date_dd','date_hh','per_last']]

#启用价格
url_address = ['https://api.glassnode.com/v1/metrics/market/price_usd_close']
url_name = ['Price']
# insert your API key here
API_KEY = '26BLocpWTcSU7sgqDdKzMHMpJDm'
data_list = []
for num in range(len(url_name)):
    addr = url_address[num]
    name = url_name[num]
    # make API request
    res_addr = requests.get(addr,params={'a': 'BTC','i':'1h', 'api_key': API_KEY})
    # convert to pandas dataframe
    ins = pd.read_json(res_addr.text, convert_dates=['t'])
    ins['date'] =  ins['t']
    ins[name] =  ins['v']
    ins = ins[['date',name]]
    data_list.append(ins)

result_data = data_list[0][['date']]
for i in range(len(data_list)):
    df = data_list[i]
    result_data = result_data.merge(df,how='left',on='date')
#last_data = result_data[(result_data.date>='2016-01-01') & (result_data.date<='2020-01-01')]
last_data = result_data[(result_data.date>='2023-05-01')]
last_data = last_data.sort_values(by=['date'])
last_data = last_data.reset_index(drop=True)
last_data['date_dd'] = last_data['date'].apply(lambda x: str(x)[0:10])
last_data['date_hh'] = last_data['date'].apply(lambda x: str(x)[11:13])
next_df = raw_data.merge(last_data,how='left',on=['date_dd','date_hh'])
next_df = next_df[next_df.date >= '2023-06-02']
next_df = next_df.reset_index(drop=True)
date = []
price = []
yijia = []
num = 4
for i in range(num,len(next_df)):
    ins = next_df[i-num:i]
    ins = ins.sort_values(by=['date'])
    ins = ins.reset_index(drop=True)
    date.append(ins['date'][num-1])
    price.append(ins['Price'][num-1])
    yijia.append(np.mean(ins['per_last']))
test_df = pd.DataFrame({'date':date,'price':price,'yijia':yijia}) 
test_df = test_df[['date','price','yijia']]
test_df_yijia = test_df[['date','yijia']]
test_df_yijia['date'] = test_df['date'] + datetime.timedelta(hours=23)
test_df_price = test_df[['date','price']]
next_df = test_df_yijia.merge(test_df_price,how='left',on=['date'])
import matplotlib.pyplot as plt
import seaborn as sns
f, axes = plt.subplots(figsize=(20, 10))
axes_fu = axes.twinx()

sns.lineplot(x="date", y="price", data=next_df, color = 'black', linewidth=0.5,ci=95, ax=axes)
sns.lineplot(x="date", y="yijia",color='red', linewidth=0.5,data=next_df, ci=95, ax=axes_fu)
plt.title('未来24小时BTC价格趋势预测', fontsize=20,fontproperties=prop) 
axes.set_xlabel('时间',fontsize=14,fontproperties=prop)
axes.set_ylabel("BTC价格",fontsize=14,fontproperties=prop)
axes_fu.set_ylabel("BTC价格趋势",fontsize=14,fontproperties=prop)
plt.savefig('future_24.png',  bbox_inches='tight')
plt.close()



# =======================================================================链上指标===================================================================

def cal(x):
    if x>= pd.to_datetime('2013-01-01') and x<= pd.to_datetime('2016-10-31'):
        y = 'Second cycle'
    elif x>= pd.to_datetime('2016-11-01') and x<= pd.to_datetime('2020-04-30'):
        y = 'Third cycle'
    else:
        y = 'Fourth cycle'
    return y
url_address = ['https://api.glassnode.com/v1/metrics/indicators/puell_multiple',
                'https://api.glassnode.com/v1/metrics/indicators/sopr_adjusted',
                'https://api.glassnode.com/v1/metrics/market/mvrv_z_score',
                'https://api.glassnode.com/v1/metrics/indicators/rhodl_ratio',
                'https://api.glassnode.com/v1/metrics/indicators/net_realized_profit_loss',
                'https://api.glassnode.com/v1/metrics/market/price_usd_close',
                'https://api.glassnode.com/v1/metrics/supply/profit_relative',
                'https://api.glassnode.com/v1/metrics/transactions/transfers_volume_to_exchanges_sum',
                'https://api.glassnode.com/v1/metrics/transactions/transfers_volume_from_exchanges_sum']
url_name = ['Puell Multiple', 'aSOPR','MVRV Z-Score','RHODL Ratio','Net Realized Profit/Loss','Price','Percent Supply in Profit','in_exchanges', 'out_exchanges']
# insert your API key here
API_KEY = '26BLocpWTcSU7sgqDdKzMHMpJDm'
data_list = []
for num in range(len(url_name)):
    print(num)
    addr = url_address[num]
    name = url_name[num]
    # make API request
    res_addr = requests.get(addr,params={'a': 'BTC', 'api_key': API_KEY})
    # convert to pandas dataframe
    ins = pd.read_json(res_addr.text, convert_dates=['t'])
    ins['date'] =  ins['t']
    ins[name] =  ins['v']
    ins = ins[['date',name]]
    data_list.append(ins)

result_data = data_list[0][['date']]
for i in range(len(data_list)):
    df = data_list[i]
    result_data = result_data.merge(df,how='left',on='date')
#last_data = result_data[(result_data.date>='2016-01-01') & (result_data.date<='2020-01-01')]
last_data = result_data[(result_data.date>='2012-10-01')]
last_data = last_data.sort_values(by=['date'])
last_data = last_data.reset_index(drop=True)
date = []
pm = []
mvrv = []
rhold = []
price = []
sopr_7 = []
sopr = []
supply = []
for j in range(len(last_data)-49):
    ins = last_data[j:j+50]
    ins = ins.reset_index(drop=True)
    date.append(ins['date'][49])
    pm.append(ins['Puell Multiple'][49])
    mvrv.append(ins['MVRV Z-Score'][49])
    rhold.append(ins['RHODL Ratio'][49])
    sopr.append(ins['aSOPR'][49])
    supply.append(ins['Percent Supply in Profit'][49])
    price.append(ins['Price'][49])
    #短期指标
    sopr_7.append(np.mean(ins['aSOPR'][-7:]))
res_df = pd.DataFrame({'date':date,'Puell Multiple':pm,'MVRV Z-Score':mvrv,'RHODL Ratio':rhold,'Price':price,'Percent Supply in Profit':supply,'7MA aSOPR':sopr_7,'aSOPR':sopr})
res_df = res_df[(res_df.date>='2019-01-01')]
res_df['cycle'] = res_df['date'].apply(lambda x:cal(x))
res_df['log(BTC price)'] = np.log(res_df['Price'])
res_df['log(RHODL Ratio)'] = np.log(res_df['RHODL Ratio'])
res_df['x1'] = 7
res_df['x2'] = 0
res_df['y1'] = 4
res_df['y2'] = 0.5
res_df['z1'] = np.log(49000)
res_df['z2'] = np.log(350)
res_df['w'] = 1
res_df['p1'] = 0.9
res_df['p2'] = 0.5

res_df['f1'] = 0.95
res_df['f2'] = 0.5

# --- asopr
f, axes = plt.subplots(figsize=(20, 10))
axes_fu = axes.twinx()
sns.lineplot(x="date", y="w", data=res_df, color='red',linestyle='--', ax=axes_fu)
sns.lineplot(x="date", y="aSOPR",color='black',data=res_df, ax=axes_fu)
sns.lineplot(x="date", y="log(BTC price)",hue = 'cycle', data=res_df, ax=axes)
axes.tick_params(labelsize=10)
axes.legend(loc='upper left', fontsize=5)


axes.set_xlabel('时间',fontsize=14,fontproperties=prop)
axes.set_ylabel("log(BTC price)",fontsize=10)
axes_fu.set_ylabel("aSOPR",fontsize=10)
plt.title('每日aSOPR值', fontsize=20,fontproperties=prop) 

plt.savefig('aSOPR.png') 
plt.close()

add_mark(file = "aSOPR.png", out = "out",mark = "MMC研究猿卡森出品", opacity=0.2, angle=30, space=30)

# --- 7ma asopr
f, axes = plt.subplots(figsize=(20, 10))
axes_fu = axes.twinx()
sns.lineplot(x="date", y="w", data=res_df, color='red',linestyle='--', ax=axes_fu)
sns.lineplot(x="date", y="7MA aSOPR",color='black',data=res_df, ax=axes_fu)
sns.lineplot(x="date", y="log(BTC price)",hue = 'cycle', data=res_df, ax=axes)
axes.tick_params(labelsize=10)
axes.legend(loc='upper left', fontsize=5)
plt.title('7MA aSOPR值', fontsize=10) 
axes.set_xlabel('时间',fontsize=14,fontproperties=prop)
axes.set_ylabel("log(BTC price)",fontsize=10)
axes_fu.set_ylabel("7MA aSOPR",fontsize=10)

plt.savefig('7MA_aSOPR.png') 
plt.close()

add_mark(file = "7MA_aSOPR.png", out = "out",mark = "MMC研究猿卡森出品", opacity=0.2, angle=30, space=30)

# --- mvrv
f, axes = plt.subplots(figsize=(20, 10))
axes_fu = axes.twinx()
sns.lineplot(x="date", y="x1", data=res_df, color = 'green',linestyle='--',ax=axes_fu)
sns.lineplot(x="date", y="x2", data=res_df, color='red',linestyle='--',ax=axes_fu)
sns.lineplot(x="date", y="MVRV Z-Score",color='black',data=res_df,ax=axes_fu)
sns.lineplot(x="date", y="log(BTC price)",hue = 'cycle', data=res_df,ax=axes)
axes.tick_params(labelsize=10)
plt.title('MVRV Z-Score值', fontsize=10) 
axes.legend(loc='upper left', fontsize=5)
axes.set_xlabel('时间',fontsize=14,fontproperties=prop)
axes.set_ylabel("log(BTC price)",fontsize=10)
axes_fu.set_ylabel("MVRV Z-Score",fontsize=10)

plt.savefig('MVRV_Z_Score.png') 
plt.close()
add_mark(file = "MVRV_Z_Score.png", out = "out",mark = "MMC研究猿卡森出品", opacity=0.2, angle=30, space=30)

# --- pm
f, axes = plt.subplots(figsize=(20, 10))
axes_fu = axes.twinx()
sns.lineplot(x="date", y="y1", data=res_df, color = 'green',linestyle='--', ax=axes_fu)
sns.lineplot(x="date", y="y2", data=res_df, color='red', linestyle='--',ax=axes_fu)
sns.lineplot(x="date", y="Puell Multiple",color='black',data=res_df, ax=axes_fu)
sns.lineplot(x="date", y="log(BTC price)",hue = 'cycle', data=res_df, ax=axes)
axes.tick_params(labelsize=10)
plt.title('Puell Multiple值', fontsize=10) 
axes.legend(loc='upper left', fontsize=5)

axes.set_xlabel('时间',fontsize=14,fontproperties=prop)
axes.set_ylabel("log(BTC price)",fontsize=10)
axes_fu.set_ylabel("Puell Multiple",fontsize=10)

#plt.show()
plt.savefig('Puell.png')
plt.close()
add_mark(file = "Puell.png", out = "out",mark = "MMC研究猿卡森出品", opacity=0.2, angle=30, space=30)

# --- rhold
f, axes = plt.subplots(figsize=(20, 10))
axes_fu = axes.twinx()
sns.lineplot(x="date", y="z1", data=res_df, color = 'green',linestyle='--',ax=axes_fu)
sns.lineplot(x="date", y="z2", data=res_df, color='red', linestyle='--',ax=axes_fu)
sns.lineplot(x="date", y="log(RHODL Ratio)",color='black',data=res_df, ax=axes_fu)
sns.lineplot(x="date", y="log(BTC price)",hue = 'cycle', data=res_df, ax=axes)
axes.tick_params(labelsize=10)
plt.title('log(RHODL Ratio)值', fontsize=10) 
axes.legend(loc='upper left', fontsize=5)

axes.set_xlabel('时间',fontsize=14,fontproperties=prop)
axes.set_ylabel("log(BTC price)",fontsize=10)
axes_fu.set_ylabel("log(RHODL Ratio)",fontsize=10)

#plt.show()
plt.savefig('RHODL.png')
plt.close()
add_mark(file = "RHODL.png", out = "out",mark = "MMC研究猿卡森出品", opacity=0.2, angle=30, space=30)

# --- supply

f, axes = plt.subplots(figsize=(20, 10))
axes_fu = axes.twinx()
sns.lineplot(x="date", y="f2", data=res_df, color='red',linestyle='--', ax=axes_fu)
sns.lineplot(x="date", y="f1", data=res_df, color='green',linestyle='--', ax=axes_fu)
sns.lineplot(x="date", y="Percent Supply in Profit",color='black',data=res_df, ax=axes_fu)
sns.lineplot(x="date", y="log(BTC price)",hue = 'cycle', data=res_df, ax=axes)
axes.tick_params(labelsize=10)
plt.title('Percent Supply in Profit值', fontsize=10) 
axes.legend(loc='upper left', fontsize=5)
axes.set_xlabel('时间',fontsize=14,fontproperties=prop)
axes.set_ylabel("log(BTC price)",fontsize=10)
axes_fu.set_ylabel("Percent Supply in Profit",fontsize=10)

#plt.show()
plt.savefig('Percent_Supply.png')
plt.close()
add_mark(file = "Percent_Supply.png", out = "out",mark = "MMC研究猿卡森出品", opacity=0.2, angle=30, space=30)

# ========================================================================附录=========================================================================
url_address = ['https://api.glassnode.com/v1/metrics/market/price_usd_close']
url_name = ['Price']
# insert your API key here
API_KEY = '26BLocpWTcSU7sgqDdKzMHMpJDm'
data_list = []
for num in range(len(url_name)):
    print(num)
    addr = url_address[num]
    name = url_name[num]
    # make API request
    res_addr = requests.get(addr,params={'a': 'BTC', 'api_key': API_KEY})
    # convert to pandas dataframe
    ins = pd.read_json(res_addr.text, convert_dates=['t'])
    ins['date'] =  ins['t']
    ins[name] =  ins['v']
    ins = ins[['date',name]]
    data_list.append(ins)

result_data = data_list[0][['date']]
for i in range(len(data_list)):
    df = data_list[i]
    result_data = result_data.merge(df,how='left',on='date')
#last_data = result_data[(result_data.date>='2016-01-01') & (result_data.date<='2020-01-01')]
btc_data = result_data[(result_data.date>='2019-10-01')]
btc_data = btc_data.sort_values(by=['date'])
btc_data = btc_data.reset_index(drop=True)
data_list = []
for num in range(len(url_name)):
    print(num)
    addr = url_address[num]
    name = url_name[num]
    # make API request
    res_addr = requests.get(addr,params={'a': 'ETH', 'api_key': API_KEY})
    # convert to pandas dataframe
    ins = pd.read_json(res_addr.text, convert_dates=['t'])
    ins['date'] =  ins['t']
    ins[name] =  ins['v']
    ins = ins[['date',name]]
    data_list.append(ins)

result_data = data_list[0][['date']]
for i in range(len(data_list)):
    df = data_list[i]
    result_data = result_data.merge(df,how='left',on='date')
#last_data = result_data[(result_data.date>='2016-01-01') & (result_data.date<='2020-01-01')]
eth_data = result_data[(result_data.date>='2019-01-01')]
eth_data = eth_data.sort_values(by=['date'])
eth_data = eth_data.reset_index(drop=True)
combine_data = eth_data.merge(btc_data,how='left',on=['date'])
combine_data['per'] = combine_data['Price_x']/combine_data['Price_y']


f, axes = plt.subplots(figsize=(20, 10))
axes_fu = axes.twinx()
sns.lineplot(x="date", y="per",color='red',data=combine_data, ax=axes_fu)
sns.lineplot(x="date", y="Price_y",color='black', data=combine_data, ax=axes)
axes.tick_params(labelsize=10)
plt.title('ETH Price/BTC Price', fontsize=10) 
axes.legend(loc='upper left', fontsize=5)
axes.set_xlabel('时间',fontsize=14,fontproperties=prop)
axes.set_ylabel("BTC price",fontsize=10)
axes_fu.set_ylabel("ETH Price/BTC Price",fontsize=10)

#plt.show()
plt.savefig('eth_btc.png')
plt.close()
add_mark(file = "eth_btc.png", out = "out",mark = "MMC研究猿卡森出品", opacity=0.2, angle=30, space=30)


# meme币
import json
import requests
import pandas as pd
import time
import numpy as np
import os
import re
import datetime
from requests import Request,Session
from datetime import datetime
from datetime import timedelta
import pandas as pd
from requests.exceptions import ConnectionError,Timeout,TooManyRedirects
url = 'https://pro-api.coinmarketcap.com/v2/cryptocurrency/ohlcv/historical'
headers = {'X-CMC_PRO_API_KEY':'4e4b92a2-8f26-48bf-91d2-11200a6441c2'}
session = Session()
session.headers.update(headers)
#74，doge,5994,shib,24478,pepe,25028,ordi
ids = [74,5994,24478,25028]
res_df_meme = pd.DataFrame()
for s_id in ids:
    response = session.get(url,params={'convert_id':2781,'id':s_id})
    data = json.loads(response.text)
    print(data)
    name = data['data']['symbol']
    print(name)
    lis = data['data']['quotes']
    close = []
    date = []
    for i in range(len(lis)):
        close.append(lis[i]['quote']['2781']['close'])
        date.append(lis[i]['quote']['2781']['timestamp'])
    df = pd.DataFrame({'date':date,'close':close})
    df['name'] = name
    res_df_meme = pd.concat([res_df_meme,df])
    time.sleep(2)
res_df_meme['date'] = res_df_meme['date'].apply(lambda x:str(x)[0:10])
res_df_doge = res_df_meme[res_df_meme.name=='DOGE']
res_df_doge['close'] = res_df_doge['close'].apply(lambda x:round(x,3))
res_df_doge = res_df_doge[['date','close']]
res_df_shib = res_df_meme[res_df_meme.name=='SHIB']
res_df_shib['close'] = res_df_shib['close'].apply(lambda x:round(x,9))
res_df_shib = res_df_shib[['date','close']]
res_df_pepe = res_df_meme[res_df_meme.name=='PEPE']
res_df_pepe['close'] = res_df_pepe['close'].apply(lambda x:round(x,7))
res_df_pepe = res_df_pepe[['date','close']]
res_df_ordi = res_df_meme[res_df_meme.name=='ORDI']
res_df_ordi['close'] = res_df_ordi['close'].apply(lambda x:round(x,3))
res_df_ordi = res_df_ordi[['date','close']]
res_df_all = res_df_doge.merge(res_df_shib,how='left',on=['date']).merge(res_df_pepe,how='left',on=['date']).merge(res_df_ordi,how='left',on=['date'])

'''
#山寨
url = 'https://pro-api.coinmarketcap.com/v1/cryptocurrency/listings/historical'
headers = {'X-CMC_PRO_API_KEY':'4e4b92a2-8f26-48bf-91d2-11200a6441c2'}
session = Session()
session.headers.update(headers)
date_now = str(datetime.utcnow())[0:10]
data_before_1 = str(pd.to_datetime(date_now) - timedelta(days=1))[0:10]
date_before = ['2023-01-01',data_before_1]

crypto_info = pd.DataFrame() 
for j in range(len(date_before)):
    date_j = date_before[j]
    print(j)
    logo = 0
    while logo == 0:
        response = session.get(url,params={'convert_id':2781,'date':date_j })
        data = json.loads(response.text)
        print(data['status']['error_code'])
        if data['status']['error_code'] != 0:
            logo = 0
        else:
            for i in range(len(data['data'])):
                id_num = data['data'][i]['id']
                name = data['data'][i]['name']
                symbol = data['data'][i]['symbol']
                date_added = data['data'][i]['date_added']
                tags = data['data'][i]['tags']
                circulating_supply = data['data'][i]['circulating_supply']
                total_supply = data['data'][i]['total_supply']
                cmc_rank = data['data'][i]['cmc_rank']
                last_updated = data['data'][i]['last_updated']   
                price = data['data'][i]['quote']['2781']['price']     
                market_cap = data['data'][i]['quote']['2781']['market_cap']      
                #print(id_num,id_num,name,symbol,date_added,tags,circulating_supply,total_supply,cmc_rank,last_updated,price,market_cap)
                df = pd.DataFrame({'date':date_j,'id':id_num,'name':name,'symbol':symbol,'date_added':date_added,'tags':str(tags),'circulating_supply':circulating_supply,'total_supply':total_supply,'cmc_rank':cmc_rank,'last_updated':last_updated,'price':price,'market_cap':market_cap},index=[0])
                crypto_info = pd.concat([crypto_info,df])
                
            logo = 1
            #crypto_info = crypto_info.reset_index()
    time.sleep(1)
crypto_info_now = crypto_info[crypto_info.date==data_before_1]
crypto_info_now = crypto_info_now[['symbol','price','market_cap']]
crypto_info_before = crypto_info[crypto_info.date=='2023-01-01']
crypto_info_before = crypto_info_before[['symbol','price','market_cap']]
crypto_tag =  crypto_info[crypto_info.date=='2023-06-18'][['symbol','tags']].drop_duplicates()
crypto_info_combine = crypto_info_now.merge(crypto_info_before,how='left',on=['symbol']).merge(crypto_tag,how='left',on=['symbol'])
crypto_info_combine = crypto_info_combine.fillna(0.00000000000001)
crypto_info_combine['per'] =  (crypto_info_combine['price_x']- crypto_info_combine['price_y'])/crypto_info_combine['price_y']
per_btc = (crypto_info_combine['price_x'][0] - crypto_info_combine['price_y'][0])/crypto_info_combine['price_y'][0]
crypto_info_last = crypto_info_combine[crypto_info_combine.per>per_btc]
crypto_info_last = crypto_info_last[['symbol','price_x','market_cap_x','tags']]
crypto_info_last['num'] = crypto_info_last.index
crypto_info_last
'''
# =======================================================================输出docment===================================================================
import sys
from docx import Document
from docx.shared import Inches
# 创建文档对象
document = Document()
# 设置文档标题，中文要用unicode字符串
document.add_heading(u'%s日BTC现货策略数据'%(str(date_now)[0:10]),0)
# 添加有序列表

document.add_paragraph('昨日BTC价格回顾',style = 'ListBullet')
document.add_paragraph('稳定币监控',style = 'ListBullet')
document.add_paragraph('链上大额转账监控',style = 'ListBullet')
document.add_paragraph('重要链上指标监控',style = 'ListBullet')
document.add_paragraph('黑天鹅事件监控',style = 'ListBullet')
document.add_paragraph('美国重要经济数据公布日期',style = 'ListBullet')


# -----  120/200/4y 均线
document.add_page_break()
#p = document.add_paragraph('This is a paragraph in new page.')
document.add_heading(u'1.BTC价格收盘回顾',level = 1)
jun_df = jun_df.reset_index(drop=True)
value_1 = jun_df['price_raw'][len(jun_df)-1]
value_120 = jun_df['price_ma120'][len(jun_df)-1]
value_200 = jun_df['price_ma200'][len(jun_df)-1]
value_4y = jun_df['price_ma4y'][len(jun_df)-1]

btc_zhangdiefu = round( ((price_data['close'][len(price_data)-2] - price_data['close'][len(price_data)-3])/price_data['close'][len(price_data)-3])*100,2)
btc_zhangdiefu_1 = str(btc_zhangdiefu)+'%'

document.add_paragraph('昨日BTC收盘价格为：%s'%(str(value_1)),style = 'ListBullet')
document.add_paragraph('昨日BTC涨幅为：%s'%(btc_zhangdiefu_1),style = 'ListBullet')
document.add_paragraph('昨日BTC收盘MA120价格为：%s'%(str(value_120)),style = 'ListBullet')
document.add_paragraph('昨日BTC收盘MA200价格为：%s'%(str(value_200)),style = 'ListBullet')
document.add_paragraph('昨日BTC收盘MA4Y价格为：%s'%(str(value_4y)),style = 'ListBullet')
# 添加图片，并指定宽度
document.add_picture('/root/xianhuo/out/BTC价格120D.png',width = Inches(5.25))

document.add_paragraph('BTC市值占比',style = 'ListBullet')
document.add_paragraph('牛市要启动，比特币市值占比至少要高于50%。',style = 'ListBullet')
btc_dominance = str(domain_data['value'][len(domain_data)-1]) + '%'
text = '今日BTC市值占比为：%s。'%(btc_dominance)
document.add_paragraph(text,style = 'ListBullet')
# 添加图片，并指定宽度
document.add_picture('/root/xianhuo/out/btc_dominance.png',width = Inches(5.25))

document.add_paragraph('交易所比特币净流入量',style = 'ListBullet')
jinliuru_data = jinliuru_data.reset_index(drop=True)

btc_netflow = jinliuru_data['value'][len(jinliuru_data)-1]
text = '昨日交易所比特币净流入量为：%s。'%(btc_netflow)
document.add_paragraph(text,style = 'ListBullet')
# 添加图片，并指定宽度
document.add_picture('/root/xianhuo/out/btc_netflow_price.png',width = Inches(5.25))

document.add_paragraph('近14日矿工每日平均流入交易所比特币量',style = 'ListBullet')
miner_data = miner_data.reset_index(drop=True)

miner_netflow = miner_data['miner_raw'][len(miner_data)-1]
text = '昨日矿工流入交易所比特币量为：%s。'%(miner_netflow)
document.add_paragraph(text,style = 'ListBullet')
# 添加图片，并指定宽度
document.add_picture('/root/xianhuo/out/btc_miner_v.png',width = Inches(5.25))

document.add_paragraph('近30日巨鲸每日平均流入交易所比特币量',style = 'ListBullet')
whale_data = whale_data.reset_index(drop=True)

whale_netflow = whale_data['whale_raw'][len(whale_data)-1]
text = '昨日巨鲸流入交易所比特币量为：%s。'%(whale_netflow)
document.add_paragraph(text,style = 'ListBullet')
# 添加图片，并指定宽度
document.add_picture('/root/xianhuo/out/btc_whale_v.png',width = Inches(5.25))


# -----  链上稳定币监控
document.add_page_break()
#p = document.add_paragraph('This is a paragraph in new page.')
document.add_heading(u'2.链上稳定币监控',level = 1)

document.add_paragraph('以下统计的稳定币包括USDT/USDC/BUSD/TUSD共四种',style = 'ListBullet')
document.add_picture('/root/xianhuo/out/BTC价格-稳定币发行总量.png',width = Inches(6))

document.add_paragraph('近7日稳定币总量值',style = 'ListBullet')
all_stable_coin = all_stable_coin[['date','Toal Supply']]
all_stable_coin = all_stable_coin.reset_index(drop=True)
sub_all_stable_coin = all_stable_coin[-7:]

t = document.add_table(rows=1, cols=2) # 插入表格，先将表头写好，参数：rows:行，cols:列
hdr_cells = t.rows[0].cells
hdr_cells[0].text = '时间' # 表头
hdr_cells[1].text = '数量'# 表头

for d in sub_all_stable_coin.values.tolist(): # 
    print(d)
    for date,Supply in [d]: # 读取每一行内容
        row_cells = t.add_row().cells # 读到一行就在word的表格中插入一行
        row_cells[0].text = str(date)
        row_cells[1].text = str(round(Supply,2))
        
#usdt的溢价率
document.add_paragraph('USDT场外溢价率监控',style = 'ListBullet')
#p = document.add_paragraph('This is a paragraph in new page.')

usdt_data = usdt_data.sort_values(by='date')
usdt_data = usdt_data.reset_index(drop=True)

usdt_price = usdt_data['usdt'][len(usdt_data)-1]
usd_price = usdt_data['usd'][len(usdt_data)-1]
usdt_per = usdt_data['per'][len(usdt_data)-1]
document.add_paragraph('当前USD兑CNY价格为：%s'%(usd_price),style = 'ListBullet')
document.add_paragraph('当前USDT场外兑CNY价格为：%s'%(usdt_price),style = 'ListBullet')
document.add_paragraph('当前USDT场外溢价率为：%s'%(usdt_per),style = 'ListBullet')

#usdt流入交易所
document.add_paragraph('交易所USDT净流入量',style = 'ListBullet')
u_jinliuru_data = u_jinliuru_data.reset_index(drop=True)
usdt_netflow = u_jinliuru_data['value'][len(u_jinliuru_data)-1]
text = '昨日交易所USDT净流入量为：%s。'%(usdt_netflow)
document.add_paragraph(text,style = 'ListBullet')
# 添加图片，并指定宽度
document.add_picture('/root/xianhuo/out/usdt_netflow_price.png',width = Inches(5.25))


# ----- 重要链上指标
document.add_page_break()
#p = document.add_paragraph('This is a paragraph in new page.')
document.add_heading(u'3.重要链上指标',level = 1)

document.add_paragraph('短期行情',style = 'ListBullet')

res_df = res_df.reset_index(drop=True)
asopr_v = res_df['aSOPR'][len(res_df)-1]
sopr_7v = res_df['7MA aSOPR'][len(res_df)-1]


document.add_paragraph('aSOPR',style = 'ListBullet')
document.add_paragraph('当aSOPR值大于1时，说明BTC全网持有者总体处于盈利状态，当其小于1时，说明总体处于亏损状态，在目前市场状态下是看涨信号,可以少量现货进入。',style = 'ListBullet')
document.add_paragraph('昨日收盘值为：%s'%(asopr_v),style = 'ListBullet')
document.add_picture('/root/xianhuo/out/aSOPR.png',width = Inches(5.25))


document.add_paragraph('7MA aSOPR',style = 'ListBullet')
document.add_paragraph('当7MA aSOPR值大于1时，说明BTC全网持有者总体处于盈利状态，当其小于1时，说明总体处于亏损状态，在目前市场状态下是可以开启定投。',style = 'ListBullet')
document.add_paragraph('昨日收盘值为：%s'%(sopr_7v),style = 'ListBullet')
document.add_picture('/root/xianhuo/out/7MA_aSOPR.png',width = Inches(5.25))

document.add_paragraph('BTC RISK INDEX',style = 'ListBullet')
risk_data = risk_data.reset_index(drop=True)
risk = risk_data['value'][len(risk_data)-1]
document.add_paragraph('昨日收盘值为：%s'%(risk),style = 'ListBullet')
document.add_picture('/root/xianhuo/out/risk_index.png',width = Inches(5.25))

document.add_paragraph('Entity-Adjusted STH-NUPL',style = 'ListBullet')
nupl_data = nupl_data.reset_index(drop=True)
nupl = nupl_data['value'][len(nupl_data)-1]
document.add_paragraph('昨日收盘值为：%s'%(nupl),style = 'ListBullet')
document.add_picture('/root/xianhuo/out/sth_nupl.png',width = Inches(5.25))



document.add_paragraph('周期行情',style = 'ListBullet')

pm_v = res_df['Puell Multiple'][len(res_df)-1]
mvrv_v = res_df['MVRV Z-Score'][len(res_df)-1]
supply_v = res_df['Percent Supply in Profit'][len(res_df)-1]
rhodl_v = res_df['RHODL Ratio'][len(res_df)-1]


document.add_paragraph('MVRV Z-Score',style = 'ListBullet')
document.add_paragraph('当MVRV Z-Score值大于7时，是牛顶信号，当其小于0时，是熊底信号。',style = 'ListBullet')      
document.add_paragraph('昨日收盘值为：%s'%(mvrv_v),style = 'ListBullet')
document.add_picture('/root/xianhuo/out/MVRV_Z_Score.png',width = Inches(5.25))

document.add_paragraph('Puell Multiple',style = 'ListBullet')
document.add_paragraph('当Puell Multiple值大于4时，是牛顶信号，当其小于0.5时，是熊底信号。',style = 'ListBullet')
document.add_paragraph('昨日收盘值为：%s'%(pm_v),style = 'ListBullet')
document.add_picture('/root/xianhuo/out/Puell.png',width = Inches(5.25))

document.add_paragraph('Percent Supply in Profit',style = 'ListBullet')
document.add_paragraph('当Percent Supply in Profit值大于0.95时，是牛顶信号，当其小于0.5时，是熊底信号。',style = 'ListBullet')
document.add_paragraph('昨日收盘值为：%s'%(supply_v),style = 'ListBullet')
document.add_picture('/root/xianhuo/out/Percent_Supply.png',width = Inches(5.25))

document.add_paragraph('RHODL Ratio',style = 'ListBullet')
document.add_paragraph('当RHODL Ratio值大于49000时，是牛顶信号，当其小于350时，是熊底信号。',style = 'ListBullet')
document.add_paragraph('昨日收盘值为：%s'%(rhodl_v),style = 'ListBullet')
document.add_picture('/root/xianhuo/out/RHODL.png',width = Inches(5.25))


# -----  链上大额转账监控
document.add_page_break()
#p = document.add_paragraph('This is a paragraph in new page.')
document.add_heading(u'4.链上大额转账监控',level = 1)

sub_zhuanzhang_df_1 = zhuanzhang_df[(zhuanzhang_df.crypto.isin (['BTC','ETH'])) & (zhuanzhang_df.value>1000)]
document.add_paragraph('大额BTC/ETH流入交易所是砸盘信号，大额USDT/USDC转入交易所是买盘信号。',style = 'ListBullet')
if len(sub_zhuanzhang_df_1) == 0:
    content_tr = '近6个小时没有超大额(价值大于1000万刀)BTC/ETH转入交易所。'
    document.add_paragraph(content_tr,style = 'ListBullet')
else:
    sub_zhuanzhang_df_1 = sub_zhuanzhang_df_1.reset_index(drop=True)
    for i in range(len(sub_zhuanzhang_df_1)):
        s_date = sub_zhuanzhang_df_1['date'][i]
        s_crypto = sub_zhuanzhang_df_1['crypto'][i]
        s_exchange = sub_zhuanzhang_df_1['exchange'][i]
        s_hash = sub_zhuanzhang_df_1['hash'][i]
        content_tr = '%s有大额%s转入%s交易所，交易哈希为：%s'%(s_date,s_crypto,s_exchange,s_hash)
        document.add_paragraph(content_tr,style = 'ListBullet')
sub_zhuanzhang_df_2 = zhuanzhang_df[(zhuanzhang_df.crypto.isin (['USDT','USDC'])) & (zhuanzhang_df.value>1000)]
if len(sub_zhuanzhang_df_2) == 0:
    content_tr_w = '近6个小时没有超大额(价值大于1000万刀)USDT/USDT转入交易所。'
    document.add_paragraph(content_tr_w,style = 'ListBullet')
else:
    sub_zhuanzhang_df_2 = sub_zhuanzhang_df_2.reset_index(drop=True)
    for i in range(len(sub_zhuanzhang_df_2)):
        s_date = sub_zhuanzhang_df_2['date'][i]
        s_crypto = sub_zhuanzhang_df_2['crypto'][i]
        s_exchange = sub_zhuanzhang_df_2['exchange'][i]
        s_hash = sub_zhuanzhang_df_2['hash'][i]
        content_tr = '%s有大额%s转入%s交易所，交易哈希为：%s'%(s_date,s_crypto,s_exchange,s_hash)
        document.add_paragraph(content_tr,style = 'ListBullet')     
zhuanzhang_df['date'] = zhuanzhang_df['date'].apply(lambda x:str(x))
#print(zhuanzhang_df)
t = document.add_table(rows=1, cols=5) # 插入表格，先将表头写好，参数：rows:行，cols:列
hdr_cells = t.rows[0].cells
hdr_cells[0].text = '时间' # 表头
hdr_cells[1].text = '币种'# 表头
hdr_cells[2].text = '转入交易所'# 表头
hdr_cells[3].text = '转入数量'# 表头
hdr_cells[4].text = '币价值'# 表头

for d in zhuanzhang_df.values.tolist(): # 
    print(d)
    for date,crypto,exchange,number,value,hash_ in [d]: # 读取每一行内容
        row_cells = t.add_row().cells # 读到一行就在word的表格中插入一行
        row_cells[0].text = date 
        row_cells[1].text = crypto
        row_cells[2].text = exchange 
        row_cells[3].text = str(number)
        row_cells[4].text = str(value)

        
# -----  黑天鹅事件
document.add_page_break()
#p = document.add_paragraph('This is a paragraph in new page.')
document.add_heading(u'5.黑天鹅事件监控',level = 1)

document.add_paragraph(content_us,style = 'ListBullet')
document.add_paragraph(content_gox,style = 'ListBullet')
# 添加图片，并指定宽度
document.add_picture('/root/xianhuo/out/US_MT.png',width = Inches(5.25))

document.add_paragraph('今年美国政府转移BTC时间',style = 'ListBullet')
t = document.add_table(rows=1, cols=3) # 插入表格，先将表头写好，参数：rows:行，cols:列
hdr_cells = t.rows[0].cells
hdr_cells[0].text = '时间' # 表头
hdr_cells[1].text = '余额'# 表头
hdr_cells[2].text = '前一天余额'# 表头

for d in last_data_us.values.tolist(): # 
    print(d)
    for date,v1,v2 in [d]: # 读取每一行内容
        row_cells = t.add_row().cells # 读到一行就在word的表格中插入一行
        row_cells[0].text = str(date) 
        row_cells[1].text = str(v1)
        row_cells[2].text = str(v2)


        
# -----  美国重要经济数据公布日期
document.add_page_break()
#p = document.add_paragraph('This is a paragraph in new page.')
document.add_heading(u'6.美国重要经济数据公布日期',level = 1)

document.add_paragraph('CME美联储加息预测：https://www.cmegroup.com/markets/interest-rates/cme-fedwatch-tool.html',style = 'ListBullet')
if len(sub_us_eco_df) == 0:
    document.add_paragraph('%s日无重要经济数据公布'%(str(date_now)[0:10]),style = 'ListBullet')
else:
    sub_us_eco_df = sub_us_eco_df.reset_index(drop=True)
    event_t = sub_us_eco_df['事件'][0]
    document.add_paragraph('%s日%s经济数据公布'%(str(date_now)[0:10],event_t),style = 'ListBullet')
document.add_paragraph('金十数据日历提醒：https://datacenter.jin10.com',style = 'ListBullet')    
t = document.add_table(rows=1, cols=5) # 插入表格，先将表头写好，参数：rows:行，cols:列
hdr_cells = t.rows[0].cells
hdr_cells[0].text = '事件' # 表头
hdr_cells[1].text = '重要性'# 表头
hdr_cells[2].text = '公布时间'# 表头
hdr_cells[3].text = '预测值'# 表头
hdr_cells[4].text = '前值'# 表头

for d in us_eco_df.values.tolist(): # 
    print(d)
    for event,importent,date,pre_value,before_value,a_date in [d]: # 读取每一行内容
        row_cells = t.add_row().cells # 读到一行就在word的表格中插入一行
        row_cells[0].text = event 
        row_cells[1].text = importent
        row_cells[2].text = str(date) 
        row_cells[3].text = pre_value
        row_cells[4].text = before_value 



# ===============================================================附录===================================================================
document.add_page_break()
#p = document.add_paragraph('This is a paragraph in new page.')
document.add_heading(u'附录',level = 1)

per_1 = combine_data['per'][len(combine_data)-1]
document.add_paragraph('ETH Price/BTC Price',style = 'ListBullet')
document.add_paragraph('当前该比值为：%s'%(per_1),style = 'ListBullet')
document.add_picture('/root/xianhuo/out/eth_btc.png',width = Inches(5.25))
'''
document.add_paragraph('山寨币',style = 'ListBullet')
t = document.add_table(rows=1, cols=5) # 插入表格，先将表头写好，参数：rows:行，cols:列
hdr_cells = t.rows[0].cells
hdr_cells[0].text = '排名' # 表头
hdr_cells[1].text = '币种'# 表头
hdr_cells[2].text = '价格'# 表头
hdr_cells[3].text = '市值'# 表头

for d in crypto_info_last.values.tolist(): # 
    print(d)
    for symbol,price_x,market_cap_x,tags,num in [d]: # 读取每一行内容
        row_cells = t.add_row().cells # 读到一行就在word的表格中插入一行
        row_cells[0].text = str(num) 
        row_cells[1].text = str(symbol)
        row_cells[2].text = str(price_x) 
        row_cells[3].text = str(round(market_cap_x,2))
'''
document.add_paragraph('MEME币',style = 'ListBullet')
document.add_paragraph('MEME币作为市场情绪的判断，一般其到达高峰就是下跌的开始。',style = 'ListBullet')
t = document.add_table(rows=1, cols=5) # 插入表格，先将表头写好，参数：rows:行，cols:列
hdr_cells = t.rows[0].cells
hdr_cells[0].text = '时间' # 表头
hdr_cells[1].text = 'DOGE'# 表头
hdr_cells[2].text = 'SHIB'# 表头
hdr_cells[3].text = 'PEPE'# 表头
hdr_cells[4].text = 'ORDI'# 表头

for d in res_df_all.values.tolist(): # 
    print(d)
    for date,doge,shib,pepe,ordi in [d]: # 读取每一行内容
        row_cells = t.add_row().cells # 读到一行就在word的表格中插入一行
        row_cells[0].text = str(date) 
        row_cells[1].text = str(doge)
        row_cells[2].text = str(shib) 
        row_cells[3].text = str(pepe)
        row_cells[4].text = str(ordi)

# 保存文档
document.save('btc.doc')


#======自动发邮件
data_value = str(date_now)[0:10]

# 调用smtplib模块。
import smtplib
# 从email包下的text模块引入MIMEText类，用于构建邮件正文内容。
from email.mime.text import MIMEText
# 从email包下的multipart模块引入MIMEMultipart类，用于构建邮件。
from email.mime.multipart import MIMEMultipart
# 从email包下的application模块引入MIMEApplication类，用于封装附件。
from email.mime.application import MIMEApplication
# 从email包下的header模块引入Header类，用于构建邮件头。
from email.header import Header

# 发件邮箱账号。
sender ='lee_daowei@163.com'
# 发件邮箱开启SMTP服务生成的授权码。
pwd = 'GKXGKVGTYBGRMAVE'
# 收件邮箱账号。
receiver=['lee_daowei@163.com']

# mail()函数用来构建邮件。
def mail():
    #创建一个带附件的实例。可通过attach()方法把构造的内容传入到邮件的整体内容中。
    #如果一封邮件中含有附件，那邮件中必须定义multipart/mixed类型。
    #MIMEMultipart()构造方法中的第一个参数，默认值为mixed。定义mixed实现构建一个带附件的邮件体。
    msg=MIMEMultipart('mixed')
    # 定义邮件头信息。发件人、收件人、邮件主题。
    msg['From']=sender
    msg['To']=receiver[0]
    subject=f'{data_value}MMC研究猿比特币现货报告'
    # 为msg对象的Subject属性赋值。
    # 实例化了一个Header邮件头对象，并赋值给msg对象的Subject属性。
    msg['Subject']=Header(subject,'utf-8')
    # 邮件正文内容。
    main_body='此邮件为自动邮件，请勿回复！'
    # 通过attach()方法把构造的内容传入到邮件的整体内容中。
    # 邮箱正文内容，第一个参数为内容，第二个参数为格式(plain为纯文本)，第三个参数为编码。
    msg.attach(MIMEText(main_body,'plain','utf-8'))
    # 以二进制只读模式打开word文档。
    with open('btc.doc','rb') as docs_file:
        # read()方法用于从文件读取指定的字节数，如果未给定或为负则读取所有。
        # MIMEApplication()构造方法创建附件对象。
        docs_part=MIMEApplication(docs_file.read())
        # 重命名附件为XXX.docx。
        docs_part.add_header('Content-Disposition','attachment',filename=f'{data_value}BTC现货策略数据.docx')
        # 通过attach()方法把构造的内容传入到邮件的整体内容中。
        msg.attach(docs_part)
        # 返回实例。
        return msg

# send_mail()函数用来发送邮件。
def send_mail(content):
    # qq邮箱SMTP服务器。
    mailhost='smtp.163.com'
    # 开启发信服务，这里使用的是加密传输。
    server=smtplib.SMTP_SSL(mailhost)
    server.connect(mailhost,465)
    # 登录指定邮箱。
    server.login(sender,pwd)
    # 获取邮件的内容。
    msg=content
    try:
        # 发送邮件。
        server.sendmail(sender,receiver,msg.as_string())
        print('邮件发送成功！')
    except:
        print('邮件发送失败！')
    # 退出登录。
    server.quit()

# 在 if __name__ == 'main': 下的代码只有在文件作为脚本直接执行时才会被执行，
# 而 import 到其他脚本中是不会被执行的。
if __name__=='__main__':
    send_mail(mail())


